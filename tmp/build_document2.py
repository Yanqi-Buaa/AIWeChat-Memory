# -*- coding: utf-8 -*-
"""
Complete document builder with figures, figure captions, and inline image support.
"""
import docx
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os

# Paths
target_path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
img_dir = r"C:\Users\Administrator\cow\tmp\images"

doc = Document(target_path)

# ========== Helper: Insert paragraph XML element after a paragraph ==========

def insert_element_after(ref_para, new_elem):
    """Insert new_elem (as OxmlElement) after ref_para's XML element."""
    ref_para._element.addnext(new_elem)

def make_para(style_name, text=''):
    """Create a w:p OxmlElement with style and optional text."""
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.append(pStyle)
    p.append(pPr)
    
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
    return p

def make_image_para(image_path, width_inches=5.5):
    """Create a w:p OxmlElement with centered image."""
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    
    ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)
    
    # Read image
    with open(image_path, 'rb') as f:
        image_blob = f.read()
    
    ext = os.path.splitext(image_path)[1].lower()
    content_types = {
        '.png': 'image/png',
        '.emf': 'image/x-emf',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
    }
    ct = content_types.get(ext, 'image/png')
    
    # Add image part to the document
    image_part = doc.part.get_or_add_image(image_path)
    rId = doc.part.relate_to(image_part, RT.IMAGE)
    
    # EMU dimensions
    cx = int(width_inches * 914400)
    cy = int(cx * 0.7)  # aspect ratio
    
    # Build drawing XML
    drawing_xml = (
        f'<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0" '
        f'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="1" name="Picture" descr="Figure"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:nvPicPr>'
        f'<pic:cNvPr id="0" name="Picture"/>'
        f'<pic:cNvPicPr/>'
        f'</pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip r:embed="{rId}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"/>'
        f'<a:stretch><a:fillRect/></a:stretch>'
        f'</pic:blipFill>'
        f'<pic:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{cx}" cy="{cy}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'</pic:spPr>'
        f'</pic:pic>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:inline>'
        f'</w:drawing>'
    )
    
    r_elem = OxmlElement('w:r')
    r_elem.append(parse_xml(drawing_xml))
    p.append(r_elem)
    return p


# ========== CONTENT ==========

overview_text = (
    '本项目紧密围绕航空发动机整机振动突出、耦合因素多、机理复杂等工程难题，'
    '以某高推重比涡扇发动机为应用对象，开展复杂机动飞行状态下航空发动机转子-支承-机匣系统结构动力学特性研究。'
    '项目突破了气动、机动惯性等复杂载荷作用下发动机结构动力学建模和分析技术，'
    '系统掌握了整机耦合振动机理和动载荷传递规律，探明了机匣连接结构、主轴承及支承构件界面等非线性因素对'
    '转子-支承-机匣系统动力学特性的影响机理和规律，建立了相应的整机结构动力学设计理论与方法，'
    '形成了三项核心成果：\n'
    '（1）复杂飞行条件下整机结构系统动力特性仿真与评估技术；\n'
    '（2）基于飞-发多源数据融合的整机振动特征提取与状态识别技术；\n'
    '（3）不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库。\n'
    '上述成果已成功应用于某高推重比涡扇发动机的研制、试飞及交付全过程，'
    '为整机变形评估与振动数据分析提供了有力支撑，应用效果良好。'
)

chengguo1_text = (
    '该项成果主要针对某型小涵道比涡扇发动机的研制需求，重点围绕其机动飞行环境及复杂的进排气条件，'
    '建立了结构特征等效的整机有限元模型，形成了覆盖不同整机运动状态和工作状态的振动响应仿真流程。'
    '在应用过程中，创新性地考虑了机动惯性载荷和气动载荷对整机结构变形的耦合效应，'
    '提出了支承约束力学特性的等效施加方法，突破了传统仿真未充分考虑变形引致支承特性改变的局限。\n\n'
    '依托该技术，全面开展了五大类核心仿真分析工作：整机结构建模及力学特性分析、'
    '机动惯性/气动载荷下整机变形仿真、机动惯性载荷下整机振动响应仿真、'
    '加力-喷管气动激励下振动响应特征分析，以及气流旋转激励下的振动响应分析。'
    '部分仿真分析结果如图1和图2所示。通过上述仿真与评估，掌握了复杂载荷对整机结构变形和振动响应的影响规律，主要包括：\n'
    '（1）明确了复杂飞行状态下的惯性与气动载荷主要通过改变整机变形能分布'
    '（表现为支承不同心、轴承内外环倾斜及转静间隙变化），'
    '进而改变转子支承约束力学特性，最终影响转子运动状态及整机振动响应；\n'
    '（2）揭示了机动惯性载荷作用下转子弯曲变形会产生附加旋转惯性力矩的机理，'
    '特别是对于双转子系统，横向载荷会引起转子间进动相互影响，'
    '导致频谱中出现如f1+f2等组合频率成分及多阶模态振动；\n'
    '（3）阐明了加力燃烧室随机宽频气动激励的作用机理，'
    '即承力结构发生受迫振动并对转子形成基础激励，导致支点处激励差异显著，'
    '使转子运动轨迹呈现出特殊的前端外花瓣、后端内花瓣形态；\n'
    '（4）掌握了在旋转惯性、气流旋转及脉动激励综合作用下的能量传递规律，'
    '转静交互激励会产生丰富的倍频特征及幅值波动，'
    '且振动能量在转子与承力结构间发生转移，'
    '显著增加了中介轴承失效及承力结构振动疲劳损伤的风险。\n\n'
    '该技术融合了机动惯性与气动载荷影响，相比仅考虑不平衡激励的传统方法，'
    '整机振动幅值预测误差降低36%，经试飞500余小时验证，'
    '成功识别了结构变形危险点及关键参数，提出了改进建议，'
    '有力支撑了后续型号衍生设计。'
)

chengguo1_end = (
    '本技术具有工程实用性，可为在研小涵道比涡扇发动机总体结构设计提供定量计算和评估方法参照，'
    '有助于提高涡扇发动机结构设计效率，实现整机振动响应和结构损伤有效控制，以达成结构完整性设计目标。'
)

chengguo2_text = (
    '该项成果主要应用于地面台架、高空台及外场飞行试验全过程的监视与诊断。'
    '某高推重比涡扇发动机是我国第一型完全独立自主研制的涡扇发动机，'
    '其工作包线范围、推重比处于世界领先水平。'
    '在外场试飞过程中，曾出现出厂振动合格的发动机左发整机振动幅值异常增大并逼近限值，'
    '而右发正常的现象，如图3所示。'
    '项目组将本项目的研究成果成功应用于该型发动机外场飞行振动数据分析，'
    '开展了基于升维扩息的整机振动特征提取及飞-发参数关联性分析，'
    '精准捕捉到转速、飞行速度及俯仰角是对该台份振动异常影响显著的关键参数。'
    '在此基础上，通过融合地面台架数据与装配平衡数据，准确识别了结构状态，'
    '明确指出飞行过程中整机结构状态较地面试车未发生根本性改变，'
    '振动异常增大系装配环节涡轮后支点不同心在飞行气动与机动惯性载荷作用下被放大，'
    '导致振动传递特性改变所致。据此判定整机结构处于第二等级，'
    '虽振幅增加但无明显振动损伤特征，建议后续重点关注整机振动特征及变化趋势，'
    '如图4和图5所示。'
)

chengguo2_end = (
    '该技术不仅掌握了基于状态切片的振动数据升维扩息方法和高阶特征张量构建方法，'
    '还实现了跨专业多源数据融合与影响因素解耦，'
    '探明了不同转速、气动热力状态及飞行机动载荷对整机振动响应的影响规律。'
    '特别是在出厂交付前，可有效识别出由制造、装配一致性差异导致的'
    '影响飞行安全的异常状态，并指导装配调整。'
    '目前，该技术已在几台份飞行试验及几十台份地面试车中得到应用，'
    '成功筛选并解决多项潜在问题，确保了交付产品质量的一致性，'
    '避免了因非计划下台导致的任务推迟，'
    '为试验考核与分解调整提供了关键决策支持。\n\n'
    '该技术体系具备良好的通用性与可扩展性。'
    '针对高转速、高负荷条件下的高推重比小涵道比涡扇发动机，'
    '能有效提取复杂工况下的关键结构状态特征，支撑排故分析及视情维护决策。'
    '对于结构相对稳健、载荷环境相对平稳的民用大涵道比涡扇发动机及直升机涡轴发动机，'
    '通过对特征维度进行简化与剪裁，同样适用其状态监测需求。'
)

chengguo3_text = (
    '该项成果采用短时傅里叶变换实现高频采样信号的频域转化与数据压缩，'
    '提取转速基频、模态频率及频域能量分布特征。'
    '基于成果中的整机特征提取与分析技术，利用四阶张量对发动机结构状态特征进行多维描述'
    '（架次-切片-参数-统计量）。'
    '通过对多台同型号发动机历史数据的集成，'
    '构建了\u201c不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库\u201d。'
    '该数据库基于实测数据提炼关键特征参数，直接服务于飞行载荷确定及振动基准建立，'
    '如图6所示。随着数据持续积累，为服役后的异常边界判定与精细化健康管理奠定了基础。\n\n'
    '该数据库主要应用于全寿命周期健康管理与标准体系建设：'
    '通过持续积累并提炼典型飞行状态下的振动特征参数，'
    '建立了实际服役环境下的整机振动水平基准与量化评估标准，'
    '不仅为批产发动机的制造装配质量评估与异常边界精确判定提供了数据资产，'
    '还可反馈指导复杂载荷环境下发动机的结构健康管理策略制定与改进设计，'
    '提升机群整体的运行安全性与任务可靠性。'
)

conclusion_text = (
    '本项目形成的\u201c复杂机动飞行状态下航空发动机整机动力特性评估、结构设计与监视诊断技术\u201d，'
    '从仿真评估、状态监视识别到数据基准构建，'
    '形成了一套自主可控的面向真实服役环境的航空发动机整机结构动力特性评估与状态识别体系。\n\n'
    '三项核心成果均已成功应用于某高推重比涡扇发动机的研制全过程：\n'
    '成果一（复杂飞行条件下整机结构系统动力特性仿真与评估技术）'
    '服务于研发设计与验证阶段，支撑了结构薄弱环节识别与整机动力学设计优化；\n'
    '成果二（基于飞-发多源数据融合的整机振动特征提取与状态识别技术）'
    '应用于地面台架与外场试飞全过程的监视诊断，有效识别多起潜在异常并指导装配调整；\n'
    '成果三（不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库）'
    '服务于全寿命周期健康管理，建立了整机振动水平基准与量化评估标准。\n\n'
    '综上所述，本项目成果在型号研制与生产中应用效果良好，'
    '有效提升了复杂服役环境下航空发动机整机结构动力学评估与状态诊断能力，'
    '保障了装备的研制进度与运行安全，具有广阔的工程应用前景。'
)


# ========== BUILD DOCUMENT ==========

print("开始构建文档...")

# === Step 1: 概述 [30] ===
p30 = doc.paragraphs[30]
p30.clear()
p30.add_run(overview_text)
last = p30
print("  [30] 概述 - Done")

# === Step 2: 成果1 [32] ===
p32 = doc.paragraphs[32]
p32.clear()
p32.add_run(chengguo1_text)

# After p32: [图1_caption] [图1_img] [图2_caption] [图2_img] [正文结尾]
elem = p32._element
for new_elem in [
    make_para('0-图题', '图1 横向惯性载荷作用下整机承力结构变形引起支点不同心'),
    make_image_para(os.path.join(img_dir, 'image2.emf'), 5.5),
    make_para('0-图题', '图2 整机结构系统弹性线示意图'),
    make_image_para(os.path.join(img_dir, 'image1.png'), 5.5),
    make_para('0-报告正文', chengguo1_end),
]:
    elem.addnext(new_elem)
    elem = new_elem
print("  [32] 成果1 + 图1+图2 - Done")

# === Step 3: 成果2 [34] ===
p34 = doc.paragraphs[34]
p34.clear()
p34.add_run(chengguo2_text)

elem = p34._element
for new_elem in [
    make_para('0-图题', '图3 某高推重比涡扇发动机在试飞中出现的振动'),
    make_image_para(os.path.join(img_dir, 'image3.emf'), 5.5),
    make_para('0-图题', '图4 某高推重比涡扇发动机外场飞行振动特征提取与分析'),
    make_image_para(os.path.join(img_dir, 'image4.emf'), 5.5),
    make_para('0-图题', '图5 本项目研究成果在某高推重比涡扇发动机上的应用效果'),
    make_image_para(os.path.join(img_dir, 'image5.emf'), 5.5),
    make_para('0-报告正文', chengguo2_end),
]:
    elem.addnext(new_elem)
    elem = new_elem
print("  [34] 成果2 + 图3+图4+图5 - Done")

# === Step 4: 成果3 [36] ===
p36 = doc.paragraphs[36]
p36.clear()
p36.add_run(chengguo3_text)

elem = p36._element
for new_elem in [
    make_para('0-图题', '图6 典型振动特征参数数据库示意图'),
    make_image_para(os.path.join(img_dir, 'image6.emf'), 5.5),
]:
    elem.addnext(new_elem)
    elem = new_elem
print("  [36] 成果3 + 图6 - Done")

# === Step 5: 成果应用结论 ===
p37 = doc.paragraphs[37]
p37._element.addnext(make_para('0-报告正文', conclusion_text))
print("  [37后] 成果应用结论 - Done")

# ========== SAVE ==========
doc.save(target_path)
print("\n✅ 文档构建完成！")
print("=" * 50)
print("结构总览：")
print("  一、概述")
print("  二、成果1的应用情况及效果（含图1、图2）")
print("  三、成果2的应用情况及效果（含图3、图4、图5）")
print("  四、成果3的应用情况及效果（含图6）")
print("  五、成果应用结论")
print("=" * 50)
