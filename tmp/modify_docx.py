"""
修改文档脚本v2：保留原格式，仅改文字内容并标红
"""
import docx
from docx.shared import RGBColor, Pt
from lxml import etree
import copy

RED = RGBColor(0xFF, 0x00, 0x00)
WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

SRC = r'A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\20260429\测试大纲\11.4-项目技术指标测试报告-北航-20260512-修改版.docx'
DST = r'A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\20260429\测试大纲\11.4-项目技术指标测试报告-北航-20260512-修改版-已修改.docx'

doc = docx.Document(SRC)


def get_first_run_format(para):
    """获取段落第一个有效run的格式信息，用于后续保留字体"""
    for run in para.runs:
        if run.text.strip():
            return {
                'font_name': run.font.name,
                'font_size': run.font.size,
                'bold': run.bold,
                'italic': run.italic,
            }
    return {'font_name': None, 'font_size': None, 'bold': None, 'italic': None}


def apply_format(run, fmt):
    """将格式应用到run"""
    if fmt.get('font_name'):
        run.font.name = fmt['font_name']
    if fmt.get('font_size'):
        run.font.size = fmt['font_size']
    if fmt.get('bold') is not None:
        run.bold = fmt['bold']
    if fmt.get('italic') is not None:
        run.italic = fmt['italic']


def clear_para(para):
    """清除段落中所有w:r元素"""
    p_element = para._element
    for r in p_element.findall('{%s}r' % WML_NS):
        p_element.remove(r)


def set_para_text_red(para, text, fmt=None):
    """替换段落文字为红色，保留原格式"""
    if fmt is None:
        fmt = get_first_run_format(para)
    clear_para(para)
    if text:  # 如果有文字才添加run
        run = para.add_run(text)
        run.font.color.rgb = RED
        apply_format(run, fmt)
        return run
    return None


# ============================================================
# 修改1: P75 - 分析变形分布，整理这段内容
# ============================================================
p75 = doc.paragraphs[75]
fmt75 = get_first_run_format(p75)
new_p75 = (
    "惯性载荷作用下，各支点产生了不同程度的横向位移，致使转子各支点连线发生倾斜："
    "1#与2#支点连线倾斜0.3273′，2#与5#支点连线倾斜1.0455′，3#与4#支点连线倾斜0.5364′。"
    "支点不同心改变了转子-支承系统的约束状态，以下分别对低压转子和高压转子的变形特征进行分析。"
)
set_para_text_red(p75, new_p75, fmt75)

# ============================================================
# 修改2: P80 - 低压、高压不同心的等效模拟
# ============================================================
p80 = doc.paragraphs[80]
fmt80 = get_first_run_format(p80)
new_p80 = "高低压转子不同心的等效模拟方法"
set_para_text_red(p80, new_p80, fmt80)

# ============================================================
# 修改3: P83-P87 - 整理速度约束特性描述
# ============================================================

# P83: 重写为完整的支承松动速度约束特性描述
p83 = doc.paragraphs[83]
fmt83 = get_first_run_format(p83)
new_p83 = (
    "惯性载荷引起的支点不同心使支承约束呈非对称状态，支承松动条件下转子非协调涡动"
    "加剧了转子与支承结构之间的碰撞效应。支承松动时转子运动状态具有以下特征："
    "进动速度不同于自转转速，运动轨迹呈现非圆周的突变特征；转子与支承发生非弹性碰撞，"
    "约束力取决于碰撞冲量与转子运动速度，具有多频周期特征；碰撞过程中的能量传递与"
    "动量交换满足动量定理，如图14所示。"
)
set_para_text_red(p83, new_p83, fmt83)

# P84: "描述：约束能量传递和变化，满足动量定理" → 已融入P83，清空
p84 = doc.paragraphs[84]
set_para_text_red(p84, "")

# P86: "又因为，则有" → 不完整公式片段，已融入P83，清空
p86 = doc.paragraphs[86]
set_para_text_red(p86, "")

# P87: 重写建模方法段落
p87 = doc.paragraphs[87]
fmt87 = get_first_run_format(p87)
new_p87 = (
    "在建模过程中，采用本项目专题二提出的轴承-支承系统约束力学特性计算方法，"
    "对碰撞过程中的动量与能量传递进行准确数学描述，从而实现对机动飞行状态下"
    "转子运动状态的准确模拟。"
)
set_para_text_red(p87, new_p87, fmt87)

# ============================================================
# 修改4: P98-P100 - 简化数据处理两段
# ============================================================

# P98: 修改标记段落，清空
p98 = doc.paragraphs[98]
set_para_text_red(p98, "")

# P99: 简化
p99 = doc.paragraphs[99]
fmt99 = get_first_run_format(p99)
new_p99 = (
    "外场飞行试验的数据处理以频域特征相似度比对为目标，处理流程包括信号预处理、"
    "频谱分析、特征数字化和相似度计算四个步骤。"
)
set_para_text_red(p99, new_p99, fmt99)

# P100: 简化
p100 = doc.paragraphs[100]
fmt100 = get_first_run_format(p100)
new_p100 = (
    "对预处理后的加速度信号进行FFT频谱分析，试验与仿真数据统一采用Hanning窗，"
    "保证两者频率分辨率一致。"
)
set_para_text_red(p100, new_p100, fmt100)

# ============================================================
# 修改5: P113 - 加入5%频率误差判据
# ============================================================
p113 = doc.paragraphs[113]
fmt113 = get_first_run_format(p113)
new_p113 = (
    "相似度计算与指标考核。在频率成分比对中，若仿真频率与实测频率的"
    "相对误差在5%以内，则认为两者为同一频率成分。"
)
set_para_text_red(p113, new_p113, fmt113)

# ============================================================
# 修改6: P121-P124 - 重组两段
# ============================================================

# P121: 修改标记段落，清空
p121 = doc.paragraphs[121]
set_para_text_red(p121, "")

# P122: "为什么仿真结果准"
p122 = doc.paragraphs[122]
fmt122 = get_first_run_format(p122)
new_p122 = (
    "仿真结果与试验吻合较好，主要原因如下：模型通过引入惯性载荷引起的支点不同心变形，"
    "等效表征了机动飞行状态下结构变形对转子动力响应的影响；轴承间隙碰撞和连接刚度损失"
    "等因素的等效模拟，保证了仿真频谱中各频率成分的相对幅值关系与实测一致。"
)
set_para_text_red(p122, new_p122, fmt122)

# P123: 已融入P122，清空
p123 = doc.paragraphs[123]
set_para_text_red(p123, "")

# P124: 差异解释
p124 = doc.paragraphs[124]
fmt124 = get_first_run_format(p124)
new_p124 = (
    "仿真与实测频谱仍存在个别频率成分的差异。以226 Hz频率成分为例，仿真分析表明"
    "其为低压转子进动频率（139 Hz）与高压激起风扇俯仰模态频率（87 Hz）的组合频率"
    "（139+87=226 Hz），而在实际运转中，系统往往优先呈现频率最低的运动状态，"
    "组合频率分量不易被激发，因此实测频谱中未出现该频率成分。"
)
set_para_text_red(p124, new_p124, fmt124)

# ============================================================
# 保存
# ============================================================
doc.save(DST)
print(f'已保存至: {DST}')
print('所有修改已用红色字体标注，原格式保留。')
