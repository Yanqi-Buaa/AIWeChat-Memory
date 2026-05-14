# -*- coding: utf-8 -*-
"""
Fix cross-reference paragraphs: merge "如" "REF" "和" "REF" "所示" into single paragraphs.
"""
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

target = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = Document(target)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def make_mixed_para(style, segments):
    """
    Create a paragraph with mixed text runs and field code elements.
    segments: list of ('text', text_str) or ('field', field_clone) or ('field_text', text_str, field_clone, after_text)
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    p.append(pPr)
    
    for seg in segments:
        if seg[0] == 'text':
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = seg[1]
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p.append(r)
        elif seg[0] == 'field':
            # Clone the field code element (which is a complete paragraph) and extract just the field runs
            field_elem = deepcopy(seg[1])
            # The field_elem is a full <w:p> element. We need to extract just the field runs inside it.
            # Find all children of the field paragraph that are field-related
            for child in list(field_elem):
                # Copy fldChar, instrText, and field-related runs
                tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
                if tag in ['r', 'fldChar', 'instrText']:
                    p.append(deepcopy(child))
                elif tag == 'pPr':
                    pass  # skip paragraph properties (already set)
                else:
                    # Might be w:r elements in there
                    p.append(deepcopy(child))
        elif seg[0] == 'field_text':
            # text before + field + text after
            if seg[1]:
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = seg[1]
                t.set(qn('xml:space'), 'preserve')
                r.append(t)
                p.append(r)
            field_elem = deepcopy(seg[2])
            for child in list(field_elem):
                tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
                if tag in ['r', 'fldChar', 'instrText']:
                    p.append(deepcopy(child))
            if seg[3]:
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = seg[3]
                t.set(qn('xml:space'), 'preserve')
                r.append(t)
                p.append(r)
    
    return p


# Collect the REF field elements from the document
# They are currently in standalone paragraphs
def find_para_by_text(text):
    """Find FIRST paragraph containing exact text."""
    for p in doc.paragraphs:
        if text == p.text.strip():
            return p
    return None

# Find and extract REF field elements
# The REF fields are in paragraphs with just "如" or "和" that have fldChar/instrText
def extract_ref_field(para):
    """Extract the REF field runs from a paragraph element."""
    elem = para._element
    field_runs = []
    for child in list(elem):
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag == 'pPr':
            continue
        field_runs.append(deepcopy(child))
    return field_runs

# Build new paragraphs for each cross-reference position
# 成果1: "部分仿真分析结果如图1和图2所示。通过上述仿真与评估..."
# 成果2: "右发正常的现象，如图3所示。" (but this one has "如" already in prev para... let me check)

# Actually, let me check current state of the doc to understand what paragraphs exist
print("Current cross-ref paragraph structure:")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt in ['如', '和', '所示', '如所示', '所示。'] or '如' in txt[:5] and len(txt) < 20:
        # Check for field codes
        flds = p._element.findall(f'{{{ns_w}}}fldChar')
        instrs = p._element.findall(f'{{{ns_w}}}instrText')
        print(f"[{i:2d}] '{txt}' fldChar={len(flds)} instrText={len(instrs)}")

print()

# Let me collect all REF field paragraphs (those with "如" or "和" that contain field codes)
ref_paras = []
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    flds = p._element.findall(f'{{{ns_w}}}fldChar')
    instrs = p._element.findall(f'{{{ns_w}}}instrText')
    if flds and txt in ['如', '和']:
        ref_paras.append((i, p, txt))

print(f"Found {len(ref_paras)} REF field paragraphs")

# Now let's fix by replacing the problematic sequences
# For each cross-ref position, we collect the paragraphs and merge them

# Position 1: 成果1 图1+图2
# Current: [...结果如][如图1][和][和图2][所示。通过...]
# Should be: 一个paragraph with "...结果如【图1】和【图2】所示。通过..."
# Actually, let me find the text segments

# Let me just find the paragraphs and merge them inline
# Find paragraph ending with "部分仿真分析结果如"
target_para = None
for p in doc.paragraphs:
    if '部分仿真分析结果如' in p.text:
        target_para = p
        break

if target_para:
    print(f"\nFound: '{target_para.text[:60]}...'")
    # Get the next few paragraphs
    idx = list(doc.paragraphs).index(target_para)
    print(f"  idx={idx}")
    
    # Collect what follows
    next_paras = []
    for j in range(idx+1, min(idx+10, len(doc.paragraphs))):
        txt = doc.paragraphs[j].text.strip()
        flds = doc.paragraphs[j]._element.findall(f'{{{ns_w}}}fldChar')
        if txt and ('如' in txt[:5] or '和' in txt[:5] or '所示' in txt[:5]):
            next_paras.append((j, doc.paragraphs[j], txt, len(flds)))
        else:
            break
    
    print("  Following paragraphs:")
    for j, p, txt, fcnt in next_paras:
        print(f"    [{j}] '{txt}' flds={fcnt}")
    
    # Now build the merged paragraph
    # Original: "...结果如" + REF1("如"+"图1") + "和" + REF2("和"+"图2") + "所示。通过..."
    
    # Extract REF1 from the first field paragraph
    if len(next_paras) >= 1:
        ref1_para = next_paras[0][1]
        ref1_runs = []
        for child in list(ref1_para._element):
            tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
            if tag != 'pPr':
                ref1_runs.append(deepcopy(child))
    # Extract REF2 
    if len(next_paras) >= 3:
        ref2_para = next_paras[2][1]
        ref2_runs = []
        for child in list(ref2_para._element):
            tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
            if tag != 'pPr':
                ref2_runs.append(deepcopy(child))
    
    # Build merged paragraph
    merged = OxmlElement('w:p')
    # pPr with style
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), '0-报告正文')
    pPr.append(pStyle)
    merged.append(pPr)
    
    # Text: "...结果如"
    base_text = target_para.text  # e.g. "...部分仿真分析结果如"
    # Actually let me strip the trailing "如" from base_text if it ends with 如
    if base_text.endswith('如'):
        before_text = base_text[:-1]
        mid_text1 = '如'
    else:
        before_text = base_text
        mid_text1 = ''
    
    # Text run before
    r1 = OxmlElement('w:r')
    t1 = OxmlElement('w:t')
    t1.text = before_text
    t1.set(qn('xml:space'), 'preserve')
    r1.append(t1)
    merged.append(r1)
    
    # REF1 run
    if mid_text1:
        r_mid = OxmlElement('w:r')
        t_mid = OxmlElement('w:t')
        t_mid.text = mid_text1
        t_mid.set(qn('xml:space'), 'preserve')
        r_mid.append(t_mid)
        merged.append(r_mid)
    
    if len(next_paras) >= 1:
        for run in ref1_runs:
            merged.append(deepcopy(run))
    
    # Text "和"
    r2 = OxmlElement('w:r')
    t2 = OxmlElement('w:t')
    t2.text = '和'
    t2.set(qn('xml:space'), 'preserve')
    r2.append(t2)
    merged.append(r2)
    
    # REF2 run
    if len(next_paras) >= 3:
        for run in ref2_runs:
            merged.append(deepcopy(run))
    
    # "所示。通过..." from the last paragraph
    last_para = next_paras[-1][1]
    last_text = last_para.text
    r_last = OxmlElement('w:r')
    t_last = OxmlElement('w:t')
    t_last.text = last_text
    t_last.set(qn('xml:space'), 'preserve')
    r_last.append(t_last)
    merged.append(r_last)
    
    # Replace: insert merged paragraph, then delete old ones
    ref_element = target_para._element
    ref_element.addnext(merged)
    
    # Delete old paragraphs (in reverse order to avoid index issues)
    for j, p, txt, fcnt in reversed(next_paras):
        p._element.getparent().remove(p._element)
    
    # Also remove the original (it becomes redundant)
    target_para.clear()
    target_para.add_run('')  # Make it empty
    
    print("  ✅ Merged 图1+图2 cross-refs")

print("\n✅ Fix applied. Saving...")
doc.save(target)
print("Saved!")
