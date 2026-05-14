import docx

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

print(f"段落总数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

# Print all paragraphs with index, style and text
for i, p in enumerate(doc.paragraphs):
    txt = p.text
    style = p.style.name
    if txt.strip():
        print(f"[{i}] style='{style}' | {repr(txt[:200])}")
    else:
        print(f"[{i}] style='{style}' | (空)")

# Also check tables
print("\n=== 表格 ===")
for ti, table in enumerate(doc.tables):
    print(f"\nTable {ti} ({len(table.rows)} x {len(table.columns)})")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f"  [{ri},{ci}] {repr(txt[:100])}")

# Check for images
from lxml import etree
body = doc.element.body
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
         'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
         'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
         'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'}
drawings = body.findall('.//w:drawing', nsmap)
print(f"\n图片数量: {len(drawings)}")

# Check for field codes
fields = body.findall('.//w:fldChar', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
print(f"域代码数量: {len(fields)}")
