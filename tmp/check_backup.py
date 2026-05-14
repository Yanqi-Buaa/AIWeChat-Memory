import docx
path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明 - 副本.docx"
doc = docx.Document(path)
print(f"段落: {len(doc.paragraphs)}, 表格: {len(doc.tables)}")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        print(f"[{i:2d}] style='{p.style.name}' | {p.text[:100]}")
