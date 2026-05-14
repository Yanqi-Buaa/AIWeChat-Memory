import docx
from lxml import etree

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明 - 副本.docx"
doc = docx.Document(path)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

print("=== 末尾段落（含图片检测）===")
for i in range(48, 71):
    p = doc.paragraphs[i]
    elem = p._element
    
    # Check for drawing (inline images)
    drawings = elem.findall(f'.//{{{ns_wp}}}inline')
    # Check for pict (old format images)
    picts = elem.findall(f'.//{{{ns_w}}}pict')
    # Check for drawings in w namespace
    wdrawings = elem.findall(f'.//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}drawing')
    
    has_img = len(drawings) > 0 or len(picts) > 0 or len(wdrawings) > 0
    txt = p.text[:80]
    style = p.style.name
    
    img_flag = " 🖼️" if has_img else ""
    print(f"[{i:2d}] style='{style}'{img_flag} | {repr(txt)}")
    if has_img:
        print(f"       inline={len(drawings)}, pict={len(picts)}, wdrawing={len(wdrawings)}")

# Also check table 2 (the sub-figure table)
print("\n=== 表2（子图）===")
if len(doc.tables) > 1:
    table = doc.tables[1]
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f"  [{ri},{ci}] {txt}")

# Check images in body paragraphs too
print("\n=== 正文段落图片检测 ===")
for i in range(29, 48):
    p = doc.paragraphs[i]
    elem = p._element
    wdrawings = elem.findall(f'.//{{http://schemas.openxmlformats.org/wordprocessingml/2006/main}}drawing')
    if wdrawings:
        print(f"[{i}] has drawing: {p.text[:60]}")
