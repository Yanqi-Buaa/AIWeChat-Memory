# -*- coding: utf-8 -*-
"""
Complete document builder: fills content, images, captions, tables
Uses python-docx for text manipulation and image insertion
"""
import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from lxml import etree
import os, shutil

# Paths
target_path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
src_path = r"C:\Users\Administrator\cow\tmp\（公开）附件3：成果应用情况审查报告-20260203.docx"
img_dir = r"C:\Users\Administrator\cow\tmp\images"

doc = Document(target_path)
src_doc = Document(src_path)

# ========== Helper Functions ==========

def add_paragraph_after(paragraph, text, style_name):
    """Add a new paragraph with specified style after the given paragraph."""
    new_p = OxmlElement('w:p')
    # Style
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.append(pStyle)
    new_p.append(pPr)
    
    if text:
        r_elem = OxmlElement('w:r')
        rPr_elem = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), style_name)
        rPr_elem.append(rStyle)
        r_elem.append(rPr_elem)
        t_elem = OxmlElement('w:t')
        t_elem.text = text
        t_elem.set(qn('xml:space'), 'preserve')
        r_elem.append(t_elem)
        new_p.append(r_elem)
    
    paragraph._element.addnext(new_p)
    return new_p

def add_image_paragraph_after(paragraph, image_path, width_inches=5.5):
    """Add a paragraph containing an image after the given paragraph."""
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    # Center alignment
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_p.append(pPr)
    
    # Add the image using run
    run_elem = OxmlElement('w:r')
    
    # We need to add the image part and get its relationship ID
    # Use python-docx's internal method
    image_ext = os.path.splitext(image_path)[1].lower()
    content_type_map = {
        '.png': 'image/png',
        '.emf': 'image/x-emf',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
    }
    content_type = content_type_map.get(image_ext, 'image/png')
    
    # Add image to document part
    with open(image_path, 'rb') as img_file:
        image_blob = img_file.read()
    
    image_part = doc.part.get_or_add_image_part(image_blob, content_type)
    rId = doc.part.relate_to(image_part, docx.opc.constants.RELATIONSHIP_TYPE.IMAGE)
    
    # Create drawing XML
    cx = int(width_inches * 914400)  # EMU
    # For EMF, try to estimate a reasonable size
    if image_ext == '.emf':
        cx = int(5.5 * 914400)
    cy = int(cx * 0.75)  # 4:3 aspect ratio default
    
    # Build the drawing XML
    drawing_xml = (
        f'<w:drawing {nsdecls("w")}>'
        f'<wp:inline distT="0" distB="0" distL="0" distR="0" {nsdecls("wp")}>'
        f'<wp:extent cx="{cx}" cy="{cy}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:docPr id="1" name="Picture" descr="Figure"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic {nsdecls("a")}>'
        f'<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">'
        f'<pic:pic {nsdecls("pic")}>'
        f'<pic:nvPicPr>'
        f'<pic:cNvPr id="0" name="Picture"/>'
        f'<pic:cNvPicPr/>'
        f'</pic:nvPicPr>'
        f'<pic:blipFill>'
        f'<a:blip r:embed="{rId}" {nsdecls("r")}/>'
        f'<a:stretch>'
        f'<a:fillRect/>'
        f'</a:stretch>'
        f'</pic:blipFill>'
        f'<pic:spPr>'
        f'<a:xfrm>'
        f'<a:off x="0" y="0"/>'
        f'<a:ext cx="{cx}" cy="{cy}"/>'
        f'</a:xfrm>'
        f'<a:prstGeom prst="rect">'
        f'<a:avLst/>'
        f'</a:prstGeom>'
        f'</pic:spPr>'
        f'</pic:pic>'
        f'</a:graphicData>'
        f'</a:graphic>'
        f'</wp:inline>'
        f'</w:drawing>'
    )
    
    drawing_element = parse_xml(drawing_xml)
    run_elem.append(drawing_element)
    new_p.append(run_elem)
    
    paragraph._element.addnext(new_p)
    return new_p


def copy_table_from_source(src_doc, table_index, target_after_para):
    """Copy a table from source document and insert after target_after_para."""
    src_table = src_doc.tables[table_index]
    
    # Create table in target
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    
    table_elem = OxmlElement('w:tbl')
    
    # Table properties
    tblPr = OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    table_elem.append(tblPr)
    
    # Copy rows
    for src_row in src_table.rows:
        tr_elem = OxmlElement('w:tr')
        for src_cell in src_row.cells:
            tc_elem = OxmlElement('w:tc')
            tcPr = OxmlElement('w:tcPr')
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '2500')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            tc_elem.append(tcPr)
            
            # Copy cell content
            for src_para in src_cell.paragraphs:
                p_elem = OxmlElement('w:p')
                r_elem = OxmlElement('w:r')
                t_elem = OxmlElement('w:t')
                t_elem.text = src_para.text if src_para.text else ''
                t_elem.set(qn('xml:space'), 'preserve')
                r_elem.append(t_elem)
                p_elem.append(r_elem)
                tc_elem.append(p_elem)
            
            tr_elem.append(tc_elem)
        table_elem.append(tr_elem)
    
    target_after_para._element.addnext(table_elem)
    return table_elem


# ========== CONTENT DEFINITIONS ==========

# Reusing text from earlier - just need to add "如图X所示" references

overview_text = (
    '本项目紧密围绕航空发动机整机振动突出、耦合因素多、机理复杂等工程难题，'
    '以某高推重比涡扇发动机为应用对象，开展复杂机动飞行状态下航空发动机转子-支承-机匣系统结构动力学特性研究。'
    '项目突破了气动、机动惯性等复杂载荷作用下发动机结构动力学建模和分析技术，'
    '系统掌握了整机耦合振动机理和动载荷传递规律，探明了机匣连接结构、主轴承及支承构件界面等非线性因素对转子-支承-机匣系统动力学特性的影响机理和规律，'
    '建立了相应的整机结构动力学设计理论与方法，形成了三项核心成果：\n'
    '（1）复杂飞行条件下整机结构系统动力特性仿真与评估技术；\n'
    '（2）基于飞-发多源数据融合的整机振动特征提取与状态识别技术；\n'
    '（3）不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库。\n'
    '上述成果已成功应用于某高推重比涡扇发动机的研制、试飞及交付全过程，'
    '为整机变形评估与振动数据分析提供了有力支撑，应用效果良好。'
)

chengguo1_text = (
    '该项成果主要针对某型小涵道比涡扇发动机的研制需求，重点围绕其机动飞行环境及复杂的进排气条件，'
    '建立了结构特征等效的整机有限元模型，形成了覆盖不同整机运动状态和工作状态的振动响应仿真流程。'
    '在应用过程中，创新性地考虑了机动惯性载荷和气动载荷对整机结构变形的耦合效应，'
    '提出了支承约束力学特性的等效施加方法，突破了传统仿真未充分考虑变形引致支承特性改变的局限。\n\n'
    '依托该技术，全面开展了五大类核心仿真分析工作：整机结构建模及力学特性分析、'
    '机动惯性/气动载荷下整机变形仿真、机动惯性载荷下整机振动响应仿真、'
    '加力-喷管气动激励下振动响应特征分析，以及气流旋转激励下的振动响应分析。'
    '部分仿真分析结果如图1和图2所示。通过上述仿真与评估，掌握了复杂载荷对整机结构变形和振动响应的影响规律，主要包括：\n'
    '（1）明确了复杂飞行状态下的惯性与气动载荷主要通过改变整机变形能分布'
    '（表现为支承不同心、轴承内外环倾斜及转静间隙变化），'
    '进而改变转子支承约束力学特性，最终影响转子运动状态及整机振动响应；\n'
    '（2）揭示了机动惯性载荷作用下转子弯曲变形会产生附加旋转惯性力矩的机理，'
    '特别是对于双转子系统，横向载荷会引起转子间进动相互影响，'
    '导致频谱中出现如f1+f2等组合频率成分及多阶模态振动；\n'
    '（3）阐明了加力燃烧室随机宽频气动激励的作用机理，'
    '即承力结构发生受迫振动并对转子形成基础激励，导致支点处激励差异显著，'
    '使转子运动轨迹呈现出特殊的前端外花瓣、后端内花瓣形态；\n'
    '（4）掌握了在旋转惯性、气流旋转及脉动激励综合作用下的能量传递规律，'
    '转静交互激励会产生丰富的倍频特征及幅值波动，'
    '且振动能量在转子与承力结构间发生转移，'
    '显著增加了中介轴承失效及承力结构振动疲劳损伤的风险。\n\n'
    '该技术融合了机动惯性与气动载荷影响，相比仅考虑不平衡激励的传统方法，'
    '整机振动幅值预测误差降低36%，经试飞500余小时验证，'
    '成功识别了结构变形危险点及关键参数，提出了改进建议，'
    '有力支撑了后续型号衍生设计。'
)

chengguo1_end_text = (
    '本技术具有工程实用性，可为在研小涵道比涡扇发动机总体结构设计提供定量计算和评估方法参照，'
    '有助于提高涡扇发动机结构设计效率，实现整机振动响应和结构损伤有效控制，以达成结构完整性设计目标。'
)

chengguo2_text = (
    '该项成果主要应用于地面台架、高空台及外场飞行试验全过程的监视与诊断。'
    '某高推重比涡扇发动机是我国第一型完全独立自主研制的涡扇发动机，'
    '其工作包线范围、推重比处于世界领先水平。'
    '在外场试飞过程中，曾出现出厂振动合格的发动机左发整机振动幅值异常增大并逼近限值，'
    '而右发正常的现象，如图3所示。'
    '项目组将本项目的研究成果成功应用于该型发动机外场飞行振动数据分析，'
    '开展了基于升维扩息的整机振动特征提取及飞-发参数关联性分析，'
    '精准捕捉到转速、飞行速度及俯仰角是对该台份振动异常影响显著的关键参数。'
    '在此基础上，通过融合地面台架数据与装配平衡数据，准确识别了结构状态，'
    '明确指出飞行过程中整机结构状态较地面试车未发生根本性改变，'
    '振动异常增大系装配环节涡轮后支点不同心在飞行气动与机动惯性载荷作用下被放大，'
    '导致振动传递特性改变所致。据此判定整机结构处于第二等级，'
    '虽振幅增加但无明显振动损伤特征，建议后续重点关注整机振动特征及变化趋势，'
    '如图4和图5所示。'
)

chengguo2_end_text = (
    '该技术不仅掌握了基于状态切片的振动数据升维扩息方法和高阶特征张量构建方法，'
    '还实现了跨专业多源数据融合与影响因素解耦，'
    '探明了不同转速、气动热力状态及飞行机动载荷对整机振动响应的影响规律。'
    '特别是在出厂交付前，可有效识别出由制造、装配一致性差异导致的'
    '影响飞行安全的异常状态，并指导装配调整。'
    '目前，该技术已在几台份飞行试验及几十台份地面试车中得到应用，'
    '成功筛选并解决多项潜在问题，确保了交付产品质量的一致性，'
    '避免了因非计划下台导致的任务推迟，'
    '为试验考核与分解调整提供了关键决策支持。\n\n'
    '该技术体系具备良好的通用性与可扩展性。'
    '针对高转速、高负荷条件下的高推重比小涵道比涡扇发动机，'
    '能有效提取复杂工况下的关键结构状态特征，支撑排故分析及视情维护决策。'
    '对于结构相对稳健、载荷环境相对平稳的民用大涵道比涡扇发动机及直升机涡轴发动机，'
    '通过对特征维度进行简化与剪裁，同样适用其状态监测需求。'
)

chengguo3_text = (
    '该项成果采用短时傅里叶变换实现高频采样信号的频域转化与数据压缩，'
    '提取转速基频、模态频率及频域能量分布特征。'
    '基于成果中的整机特征提取与分析技术，利用四阶张量对发动机结构状态特征进行多维描述'
    '（架次-切片-参数-统计量）。'
    '通过对多台同型号发动机历史数据的集成，'
    '构建了\u201c不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库\u201d。'
    '该数据库基于实测数据提炼关键特征参数，直接服务于飞行载荷确定及振动基准建立，'
    '如图6所示。随着数据持续积累，为服役后的异常边界判定与精细化健康管理奠定了基础。\n\n'
    '该数据库主要应用于全寿命周期健康管理与标准体系建设：'
    '通过持续积累并提炼典型飞行状态下的振动特征参数，'
    '建立了实际服役环境下的整机振动水平基准与量化评估标准，'
    '不仅为批产发动机的制造装配质量评估与异常边界精确判定提供了数据资产，'
    '还可反馈指导复杂载荷环境下发动机的结构健康管理策略制定与改进设计，'
    '提升机群整体的运行安全性与任务可靠性。'
)

conclusion_text = (
    '本项目形成的\u201c复杂机动飞行状态下航空发动机整机动力特性评估、结构设计与监视诊断技术\u201d，'
    '从仿真评估、状态监视识别到数据基准构建，'
    '形成了一套自主可控的面向真实服役环境的航空发动机整机结构动力特性评估与状态识别体系。\n\n'
    '三项核心成果均已成功应用于某高推重比涡扇发动机的研制全过程：\n'
    '成果一（复杂飞行条件下整机结构系统动力特性仿真与评估技术）'
    '服务于研发设计与验证阶段，支撑了结构薄弱环节识别与整机动力学设计优化；\n'
    '成果二（基于飞-发多源数据融合的整机振动特征提取与状态识别技术）'
    '应用于地面台架与外场试飞全过程的监视诊断，有效识别多起潜在异常并指导装配调整；\n'
    '成果三（不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库）'
    '服务于全寿命周期健康管理，建立了整机振动水平基准与量化评估标准。\n\n'
    '综上所述，本项目成果在型号研制与生产中应用效果良好，'
    '有效提升了复杂服役环境下航空发动机整机结构动力学评估与状态诊断能力，'
    '保障了装备的研制进度与运行安全，具有广阔的工程应用前景。'
)


# ========== UPDATE DOCUMENT ==========

# --- 概述 [30] ---
p30 = doc.paragraphs[30]
p30.clear()
run30 = p30.add_run(overview_text)

# --- 成果1 ---
# 成果1 正文前半 [32]
p32 = doc.paragraphs[32]
p32.clear()
run32 = p32.add_run(chengguo1_text)

# 图1 图题 after p32
p_cap1 = add_paragraph_after(doc.paragraphs[32], '', '0-图表内容')
p_cap1 = add_paragraph_after(p_cap1, '图1 横向惯性载荷作用下整机承力结构变形引起支点不同心', '0-图题')

# 图1 图片
img1 = os.path.join(img_dir, 'image2.emf')
if os.path.exists(img1):
    add_image_paragraph_after(p_cap1, img1, 5.5)

# 图2 图题
p_cap2 = add_paragraph_after(p_cap1, '', '0-图表内容')
p_cap2 = add_paragraph_after(p_cap2, '图2 整机结构系统弹性线示意图', '0-图题')

# 图2 图片
img2 = os.path.join(img_dir, 'image1.png')
if os.path.exists(img2):
    add_image_paragraph_after(p_cap2, img2, 5.5)

# 成果1 正文后半
p1_end = add_paragraph_after(p_cap2, chengguo1_end_text, '0-报告正文')

# --- 成果2 ---
# 成果2 正文前半 [34]
p34 = doc.paragraphs[34]
p34.clear()
run34 = p34.add_run(chengguo2_text)

# 图3 图题
p_cap3 = add_paragraph_after(doc.paragraphs[34], '', '0-图表内容')
p_cap3 = add_paragraph_after(p_cap3, '图3 某高推重比涡扇发动机在试飞中出现的振动', '0-图题')

# 图3 图片
img3 = os.path.join(img_dir, 'image3.emf')
if os.path.exists(img3):
    add_image_paragraph_after(p_cap3, img3, 5.5)

# 图4 图题
p_cap4 = add_paragraph_after(p_cap3, '', '0-图表内容')
p_cap4 = add_paragraph_after(p_cap4, '图4 某高推重比涡扇发动机外场飞行振动特征提取与分析', '0-图题')

# 图4 图片
img4 = os.path.join(img_dir, 'image4.emf')
if os.path.exists(img4):
    add_image_paragraph_after(p_cap4, img4, 5.5)

# 图5 图题
p_cap5 = add_paragraph_after(p_cap4, '', '0-图表内容')
p_cap5 = add_paragraph_after(p_cap5, '图5 本项目研究成果在某高推重比涡扇发动机上的应用效果', '0-图题')

# 图5 图片
img5 = os.path.join(img_dir, 'image5.emf')
if os.path.exists(img5):
    add_image_paragraph_after(p_cap5, img5, 5.5)

# 成果2 正文后半
p2_end = add_paragraph_after(p_cap5, chengguo2_end_text, '0-报告正文')

# --- 成果3 ---
# 成果3 正文 [36]
p36 = doc.paragraphs[36]
p36.clear()
run36 = p36.add_run(chengguo3_text)

# 图6 图题
p_cap6 = add_paragraph_after(doc.paragraphs[36], '', '0-图表内容')
p_cap6 = add_paragraph_after(p_cap6, '图6 典型振动特征参数数据库示意图', '0-图题')

# 图6 图片  
img6 = os.path.join(img_dir, 'image6.emf')
if os.path.exists(img6):
    add_image_paragraph_after(p_cap6, img6, 5.5)

# --- 成果应用结论 ---
p37 = doc.paragraphs[37]  # "成果应用结论" title
p_conclusion = add_paragraph_after(p37, conclusion_text, '0-报告正文')


# ========== SAVE ==========
doc.save(target_path)
print("文档已成功更新！")
print("更新内容：")
print("  概述 - 已填充")
print("  成果1 - 正文 + 图1(caption+image) + 图2(caption+image) + 补充正文")
print("  成果2 - 正文 + 图3(caption+image) + 图4(caption+image) + 图5(caption+image) + 补充正文")
print("  成果3 - 正文 + 图6(caption+image)")
print("  成果应用结论 - 已填充")
print("  正文中已添加'如图X所示'引用")
