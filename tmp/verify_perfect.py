import docx
from lxml import etree

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

print(f"段落总数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

# Check key paragraphs
print("\n=== 正文结构 ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    style = p.style.name
    if txt and (any(k in txt[:15] for k in ['概述', '成果', '结论', '本项', '该项', '上述',
               '图1 ', '图2 ', '图3 ', '图4 ', '图5 ', '图6 ',
               '（1）', '（2）', '（3）', '（4）', '三项', '综上',
               '依托', '部分', '项目组', '在外', '该技术融合', '该技术不仅',
               '该技术体系', '该数据库', '此基础上', '据此'])):
        # Check for images
        elem = p._element
        imgs = elem.findall(f'.//{{{ns_wp}}}inline')
        img_flag = ' 🖼️' if imgs else ''
        print(f"[{i:2d}] {style:15s}{img_flag} | {txt[:110]}")

# Count images
body = doc.element.body
all_images = body.findall(f'.//{{{ns_wp}}}inline')
print(f"\n图片总数: {len(all_images)}")

# Count SEQ fields
seq_fields = body.findall(f'.//{{{ns_w}}}instrText')
seq_count = sum(1 for s in seq_fields if s.text and 'SEQ' in s.text)
ref_count = sum(1 for s in seq_fields if s.text and 'REF' in s.text)
print(f"SEQ域代码: {seq_count}, REF域代码: {ref_count}")

# Check table 0 unchanged
print("\n=== 第1页表格 ===")
for ri, row in enumerate(doc.tables[0].rows):
    cells = [cell.text.strip()[:60] for cell in row.cells]
    print(f"  Row {ri}: {' | '.join(cells)}")

# Check for remaining end content
end_left = sum(1 for p in doc.paragraphs if '交叉引用' in p.text)
print(f"\n末尾'交叉引用'残留: {end_left} (应为0)")

# Check if FIG markers are gone
markers = sum(1 for p in doc.paragraphs if '【XREF' in p.text or '【FIG' in p.text)
print(f"标记残留: {markers} (应为0)")

# Word count
total = sum(len(p.text) for p in doc.paragraphs)
print(f"\n文档总字数: {total}")
print("\n✅ 验证完成")
