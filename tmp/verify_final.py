import docx

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

print(f"段落总数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

# Count images
from lxml import etree
nsmap = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
}
body = doc.element.body
drawings = body.findall('.//w:drawing', nsmap)
print(f"图片数量: {len(drawings)}")

# Check key paragraphs
print("\n=== 关键段落检查 ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt and ('图' in txt[:10] or '概述' in txt[:5] or '成果' in txt[:5] or '结论' in txt[:5] or '本项' in txt[:5]):
        style = p.style.name
        print(f"[{i}] style='{style}' | {txt[:100]}...")

# Check page 1 table unchanged
print("\n=== 第1页表格 ===")
for ri, row in enumerate(doc.tables[0].rows):
    cells = [cell.text.strip()[:40] for cell in row.cells]
    print(f"  Row {ri}: {' | '.join(cells)}")

# Count figure captions and "如图" references
fig_refs = sum(1 for p in doc.paragraphs if '如图' in p.text)
captions = sum(1 for p in doc.paragraphs if p.text.strip().startswith('图') and '图' in p.text[:5])
print(f"\n'如图X所示'引用: {fig_refs}处")
print(f"图题(图X): {captions}个")

# Total body word count
total = sum(len(p.text) for p in doc.paragraphs)
print(f"文档总字数: {total}字")
