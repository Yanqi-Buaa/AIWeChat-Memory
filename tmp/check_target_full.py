import docx

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

print("=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    print(f"[{i}] style='{p.style.name}' runs={len(p.runs)} | full_text={repr(p.text)}")

print("\n=== TABLES ===")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ({len(table.rows)} x {len(table.columns)}) ---")
    for ri, row in enumerate(table.rows):
        cells = []
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip().replace('\n', '\\n')
            cells.append(f"[{ci}]{txt[:100]}")
        print(f"  Row {ri}: {' | '.join(cells)}")
