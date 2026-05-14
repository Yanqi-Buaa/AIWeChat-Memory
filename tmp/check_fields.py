import docx
from lxml import etree

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Examine field codes in paragraphs [48]-[70] (the end section)
print("=== 文件末尾段落域代码检查 ===")
for i in range(48, 71):
    p = doc.paragraphs[i]
    if i < len(doc.paragraphs):
        p_elem = p._element
        fld_chars = p_elem.findall(f'.//{{{ns_w}}}fldChar')
        instr_texts = p_elem.findall(f'.//{{{ns_w}}}instrText')
        
        has_field = len(fld_chars) > 0 or len(instr_texts) > 0
        txt = p.text.strip()
        
        if has_field or txt:
            print(f"\n[{i}] style='{p.style.name}' text='{txt[:80]}'")
            for fld in fld_chars:
                fld_type = fld.get(f'{{{ns_w}}}fldCharType')
                print(f"    fldChar: {fld_type}")
            for instr in instr_texts:
                print(f"    instrText: {instr.text}")

# Check all field codes in document
print("\n\n=== 全文域代码统计 ===")
all_fields = doc.element.body.findall(f'.//{{{ns_w}}}fldChar')
print(f"fldChar 总数: {len(all_fields)}")

# Check what SEQ fields exist
all_instr = doc.element.body.findall(f'.//{{{ns_w}}}instrText')
seq_fields = []
for instr in all_instr:
    if instr.text and 'SEQ' in instr.text:
        seq_fields.append(instr.text)
        print(f"SEQ: {instr.text}")

print(f"\nSEQ域代码数: {len(seq_fields)}")

# Check REF fields (cross references)
ref_fields = []
for instr in all_instr:
    if instr.text and 'REF' in instr.text:
        ref_fields.append(instr.text)
        print(f"REF: {instr.text}")

print(f"\nREF域代码数: {len(ref_fields)}")

# Check what bookmarks exist
bookmarks = doc.element.body.findall(f'.//{{{ns_w}}}bookmarkStart')
print(f"\n书签数量: {len(bookmarks)}")
for bm in bookmarks:
    name = bm.get(f'{{{ns_w}}}name')
    print(f"  书签: {name}")
