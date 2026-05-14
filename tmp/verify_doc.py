import docx

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

print("=== 最终文档验证 ===")
print(f"段落总数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")
print()

for i in [0,1,2,23,29,30,31,32,33,34,35,36,37,38]:
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        style = p.style.name
        txt = p.text
        if txt.strip():
            print(f"[{i}] style='{style}' | {txt[:150]}...")
            if len(txt) > 150:
                print(f"     ...(共{len(txt)}字符)")
        else:
            print(f"[{i}] style='{style}' | (空)")

# Check total length of each section
print("\n=== 各节字数统计 ===")
sections = {
    "概述": 30,
    "成果1": 32,
    "成果2": 34,
    "成果3": 36,
    "成果应用结论": 38  # the new paragraph after [37]
}
for name, idx in sections.items():
    if idx < len(doc.paragraphs):
        txt = doc.paragraphs[idx].text
        print(f"  {name}: {len(txt)} 字")
        
# Count total body text
total = sum(len(doc.paragraphs[i].text) for i in [30,32,34,36,38])
print(f"\n  正文总字数: {total} 字")

# Check table not changed
print("\n=== 第1页表格验证 ===")
for ri, row in enumerate(doc.tables[0].rows):
    cells = [cell.text.strip()[:60] for cell in row.cells]
    print(f"  Row {ri}: {' | '.join(cells)}")
