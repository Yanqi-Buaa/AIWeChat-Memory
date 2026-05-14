import docx
import sys

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
print(f"Opening: {path}")
doc = docx.Document(path)
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
for i, p in enumerate(doc.paragraphs):
    print(f"[{i}] style={p.style.name} | runs={len(p.runs)} | text={repr(p.text[:100])}")
