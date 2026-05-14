import docx, os

output_dir = r"C:\Users\Administrator\cow\tmp"

files = [
    "（公开）附件1：应用证明需求申请表-20260203.docx",
    "（公开）附件3：成果应用情况审查报告-20260203.docx",
    "（公开）附件4：应用证明模板-20260203.docx"
]

for fname in files:
    path = os.path.join(output_dir, fname)
    print(f"\n{'='*70}")
    print(f"FILE: {fname}")
    print('='*70)
    try:
        doc = docx.Document(path)
        print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
        for i, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            if txt:
                print(f"[{i}] {p.style.name} | {txt}")
        if doc.tables:
            for ti, table in enumerate(doc.tables):
                print(f"\n--- Table {ti} ({len(table.rows)} x {len(table.columns)}) ---")
                for ri, row in enumerate(table.rows):
                    cells = [cell.text.strip()[:80] for cell in row.cells]
                    print(f"  Row {ri}: {' | '.join(cells)}")
    except Exception as e:
        print(f"ERROR: {e}")
