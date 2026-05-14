# -*- coding: utf-8 -*-
"""
Clean rebuild v2: Write full paragraphs with placeholders, then inline-replace with REF fields.
"""
import shutil
from copy import deepcopy
from lxml import etree
from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import OxmlElement, parse_xml

backup = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明 - 副本.docx"
target = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"

shutil.copy2(backup, target)
doc = Document(target)
ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# ====== Collect elements from end section ======
# Cross-ref REF field elements (full paragraph elements)
xref_paras = [doc.paragraphs[i]._element for i in range(49, 55)]

# Figure image+caption elements
fig_elems = {
    1: [doc.paragraphs[55]._element, doc.paragraphs[56]._element],
    2: [doc.paragraphs[58]._element, doc.paragraphs[59]._element],
    3: [doc.paragraphs[61]._element, doc.paragraphs[62]._element],
    4: [doc.paragraphs[63]._element, doc.paragraphs[64]._element],
    5: [doc.paragraphs[65]._element],
    6: [doc.paragraphs[68]._element, doc.paragraphs[69]._element],
}

# Sub-figure table
table1 = None
if len(doc.tables) > 1:
    table1 = doc.tables[1]._tbl

print(f"Collected: {len(xref_paras)} REF fields, 6 figure sets, {'table' if table1 else 'no'} sub-table")

# ====== Helper: create a paragraph with placeholder markers ======
def make_p(style, text):
    """Create a w:p element with style and text."""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style)
    pPr.append(pStyle)
    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
    return p

def make_h1(text):
    return make_p('0-标题1', text)

def make_body(text):
    return make_p('0-报告正文', text)

# ====== Helper: inline replacement of marker with REF field ======
def replace_marker_in_para(para, marker, ref_para_elem):
    """Replace marker text within a paragraph with the REF field elements from ref_para_elem."""
    p_elem = para._element
    full_text = para.text
    
    if marker not in full_text:
        return False
    
    # Get the text before and after marker
    parts = full_text.split(marker, 1)
    before_text = parts[0]
    after_text = parts[1] if len(parts) > 1 else ''
    
    # Clear the paragraph
    for child in list(p_elem):
        if child.tag == f'{{{ns_w}}}pPr':
            continue
        p_elem.remove(child)
    
    # Add text run for "before" content
    if before_text:
        r1 = OxmlElement('w:r')
        t1 = OxmlElement('w:t')
        t1.text = before_text
        t1.set(qn('xml:space'), 'preserve')
        r1.append(t1)
        p_elem.append(r1)
    
    # Extract field runs from the REF paragraph (skip pPr)
    for child in list(ref_para_elem):
        tag = child.tag.split('}')[1] if '}' in child.tag else child.tag
        if tag != 'pPr':
            p_elem.append(deepcopy(child))
    
    # Add text run for "after" content
    if after_text:
        r2 = OxmlElement('w:r')
        t2 = OxmlElement('w:t')
        t2.text = after_text
        t2.set(qn('xml:space'), 'preserve')
        r2.append(t2)
        p_elem.append(r2)
    
    return True

# ====== Build body content (as a list of elements) ======
print("Building body content...")
body_items = []  # list of (type, data) where type='elem' or 'marker_action'

# --- 概述 ---
body_items.append(('elem', make_body(
    '本项目紧密围绕航空发动机整机振动突出、耦合因素多、机理复杂等工程难题，'
    '以某高推重比涡扇发动机为应用对象，开展复杂机动飞行状态下航空发动机转子-支承-机匣系统结构动力学特性研究。'
    '项目突破了气动、机动惯性等复杂载荷作用下发动机结构动力学建模和分析技术，'
    '系统掌握了整机耦合振动机理和动载荷传递规律，探明了机匣连接结构、主轴承及支承构件界面等非线性因素对'
    '转子-支承-机匣系统动力学特性的影响机理和规律，建立了相应的整机结构动力学设计理论与方法，'
    '形成了三项核心成果：')))
body_items.append(('elem', make_body('（1）复杂飞行条件下整机结构系统动力特性仿真与评估技术；')))
body_items.append(('elem', make_body('（2）基于飞-发多源数据融合的整机振动特征提取与状态识别技术；')))
body_items.append(('elem', make_body('（3）不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库。')))
body_items.append(('elem', make_body(
    '上述成果已成功应用于某高推重比涡扇发动机的研制、试飞及交付全过程，'
    '为整机变形评估与振动数据分析提供了有力支撑，应用效果良好。')))
print("  概述")

# --- 成果1 ---
body_items.append(('elem', make_h1('成果1的应用情况及效果')))
body_items.append(('elem', make_body(
    '该项成果主要针对某型小涵道比涡扇发动机的研制需求，重点围绕其机动飞行环境及复杂的进排气条件，'
    '建立了结构特征等效的整机有限元模型，形成了覆盖不同整机运动状态和工作状态的振动响应仿真流程。'
    '在应用过程中，创新性地考虑了机动惯性载荷和气动载荷对整机结构变形的耦合效应，'
    '提出了支承约束力学特性的等效施加方法，突破了传统仿真未充分考虑变形引致支承特性改变的局限。')))
body_items.append(('elem', make_body(
    '依托该技术，全面开展了五大类核心仿真分析工作：整机结构建模及力学特性分析、'
    '机动惯性/气动载荷下整机变形仿真、机动惯性载荷下整机振动响应仿真、'
    '加力-喷管气动激励下振动响应特征分析，以及气流旋转激励下的振动响应分析。'
    '部分仿真分析结果如【REF1】和【REF2】所示。通过上述仿真与评估，'
    '掌握了复杂载荷对整机结构变形和振动响应的影响规律，主要包括：')))
body_items.append(('elem', make_body(
    '（1）明确了复杂飞行状态下的惯性与气动载荷主要通过改变整机变形能分布'
    '（表现为支承不同心、轴承内外环倾斜及转静间隙变化），'
    '进而改变转子支承约束力学特性，最终影响转子运动状态及整机振动响应；')))
body_items.append(('elem', make_body(
    '（2）揭示了机动惯性载荷作用下转子弯曲变形会产生附加旋转惯性力矩的机理，'
    '特别是对于双转子系统，横向载荷会引起转子间进动相互影响，'
    '导致频谱中出现如f1+f2等组合频率成分及多阶模态振动；')))
body_items.append(('elem', make_body(
    '（3）阐明了加力燃烧室随机宽频气动激励的作用机理，'
    '即承力结构发生受迫振动并对转子形成基础激励，导致支点处激励差异显著，'
    '使转子运动轨迹呈现出特殊的前端外花瓣、后端内花瓣形态；')))
body_items.append(('elem', make_body(
    '（4）掌握了在旋转惯性、气流旋转及脉动激励综合作用下的能量传递规律，'
    '转静交互激励会产生丰富的倍频特征及幅值波动，'
    '且振动能量在转子与承力结构间发生转移，'
    '显著增加了中介轴承失效及承力结构振动疲劳损伤的风险。')))
body_items.append(('elem', make_body(
    '该技术融合了机动惯性与气动载荷影响，相比仅考虑不平衡激励的传统方法，'
    '整机振动幅值预测误差降低36%，经试飞500余小时验证，'
    '成功识别了结构变形危险点及关键参数，提出了改进建议，'
    '有力支撑了后续型号衍生设计。')))
# Fig1 + Fig2
body_items.append(('fig', (fig_elems[1], fig_elems[2])))
body_items.append(('elem', make_body(
    '本技术具有工程实用性，可为在研小涵道比涡扇发动机总体结构设计提供定量计算和评估方法参照，'
    '有助于提高涡扇发动机结构设计效率，实现整机振动响应和结构损伤有效控制，以达成结构完整性设计目标。')))
print("  成果1 + 图1+图2")

# --- 成果2 ---
body_items.append(('elem', make_h1('成果2的应用情况及效果')))
body_items.append(('elem', make_body(
    '该项成果主要应用于地面台架、高空台及外场飞行试验全过程的监视与诊断。'
    '某高推重比涡扇发动机是我国第一型完全独立自主研制的涡扇发动机，'
    '其工作包线范围、推重比处于世界领先水平。'
    '在外场试飞过程中，曾出现出厂振动合格的发动机左发整机振动幅值异常增大并逼近限值，'
    '而右发正常的现象，如【REF3】所示。')))
body_items.append(('elem', make_body(
    '项目组将本项目的研究成果成功应用于该型发动机外场飞行振动数据分析，'
    '开展了基于升维扩息的整机振动特征提取及飞-发参数关联性分析，'
    '精准捕捉到转速、飞行速度及俯仰角是对该台份振动异常影响显著的关键参数。'
    '在此基础上，通过融合地面台架数据与装配平衡数据，准确识别了结构状态，'
    '明确指出飞行过程中整机结构状态较地面试车未发生根本性改变，'
    '振动异常增大系装配环节涡轮后支点不同心在飞行气动与机动惯性载荷作用下被放大，'
    '导致振动传递特性改变所致。据此判定整机结构处于第二等级，'
    '虽振幅增加但无明显振动损伤特征，建议后续重点关注整机振动特征及变化趋势，'
    '如【REF4】和【REF5】所示。')))
body_items.append(('fig', (fig_elems[3], fig_elems[4], fig_elems[5])))
body_items.append(('table', table1))
body_items.append(('elem', make_body(
    '该技术不仅掌握了基于状态切片的振动数据升维扩息方法和高阶特征张量构建方法，'
    '还实现了跨专业多源数据融合与影响因素解耦，'
    '探明了不同转速、气动热力状态及飞行机动载荷对整机振动响应的影响规律。'
    '特别是在出厂交付前，可有效识别出由制造、装配一致性差异导致的'
    '影响飞行安全的异常状态，并指导装配调整。'
    '目前，该技术已在几台份飞行试验及几十台份地面试车中得到应用，'
    '成功筛选并解决多项潜在问题，确保了交付产品质量的一致性，'
    '避免了因非计划下台导致的任务推迟，'
    '为试验考核与分解调整提供了关键决策支持。')))
body_items.append(('elem', make_body(
    '该技术体系具备良好的通用性与可扩展性。'
    '针对高转速、高负荷条件下的高推重比小涵道比涡扇发动机，'
    '能有效提取复杂工况下的关键结构状态特征，支撑排故分析及视情维护决策。'
    '对于结构相对稳健、载荷环境相对平稳的民用大涵道比涡扇发动机及直升机涡轴发动机，'
    '通过对特征维度进行简化与剪裁，同样适用其状态监测需求。')))
print("  成果2 + 图3+图4+图5")

# --- 成果3 ---
body_items.append(('elem', make_h1('成果3的应用情况及效果')))
body_items.append(('elem', make_body(
    '该项成果采用短时傅里叶变换实现高频采样信号的频域转化与数据压缩，'
    '提取转速基频、模态频率及频域能量分布特征。'
    '基于成果中的整机特征提取与分析技术，利用四阶张量对发动机结构状态特征进行多维描述'
    '（架次-切片-参数-统计量）。'
    '通过对多台同型号发动机历史数据的集成，'
    '构建了\u201c不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库\u201d。'
    '该数据库基于实测数据提炼关键特征参数，直接服务于飞行载荷确定及振动基准建立，'
    '如【REF6】所示。随着数据持续积累，为服役后的异常边界判定与精细化健康管理奠定了基础。')))
body_items.append(('fig', (fig_elems[6],)))
body_items.append(('elem', make_body(
    '该数据库主要应用于全寿命周期健康管理与标准体系建设：'
    '通过持续积累并提炼典型飞行状态下的振动特征参数，'
    '建立了实际服役环境下的整机振动水平基准与量化评估标准，'
    '不仅为批产发动机的制造装配质量评估与异常边界精确判定提供了数据资产，'
    '还可反馈指导复杂载荷环境下发动机的结构健康管理策略制定与改进设计，'
    '提升机群整体的运行安全性与任务可靠性。')))
print("  成果3 + 图6")

# --- 结论 ---
body_items.append(('elem', make_h1('成果应用结论')))
body_items.append(('elem', make_body(
    '本项目形成的\u201c复杂机动飞行状态下航空发动机整机动力特性评估、结构设计与监视诊断技术\u201d，'
    '从仿真评估、状态监视识别到数据基准构建，'
    '形成了一套自主可控的面向真实服役环境的航空发动机整机结构动力特性评估与状态识别体系。')))
body_items.append(('elem', make_body(
    '三项核心成果均已成功应用于某高推重比涡扇发动机的研制全过程：')))
body_items.append(('elem', make_body(
    '成果一（复杂飞行条件下整机结构系统动力特性仿真与评估技术）'
    '服务于研发设计与验证阶段，支撑了结构薄弱环节识别与整机动力学设计优化；')))
body_items.append(('elem', make_body(
    '成果二（基于飞-发多源数据融合的整机振动特征提取与状态识别技术）'
    '应用于地面台架与外场试飞全过程的监视诊断，有效识别多起潜在异常并指导装配调整；')))
body_items.append(('elem', make_body(
    '成果三（不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库）'
    '服务于全寿命周期健康管理，建立了整机振动水平基准与量化评估标准。')))
body_items.append(('elem', make_body(
    '综上所述，本项目成果在型号研制与生产中应用效果良好，'
    '有效提升了复杂服役环境下航空发动机整机结构动力学评估与状态诊断能力，'
    '保障了装备的研制进度与运行安全，具有广阔的工程应用前景。')))
print("  结论")

# ====== Insert into document ======
print("\nInserting content...")
body = doc.element.body
# Keep paragraphs [0]-[28], remove everything after
p28_elem = doc.paragraphs[28]._element

# Collect elements to remove (from [29] to end)
to_remove = []
next_elem = doc.paragraphs[29]._element
while next_elem is not None:
    to_remove.append(next_elem)
    next_elem = next_elem.getnext()

for elem in to_remove:
    body.remove(elem)

# Now insert new content
current = p28_elem
for item_type, data in body_items:
    if item_type == 'elem':
        elem = data  # already an OxmlElement
        current.addnext(elem)
        current = elem
    elif item_type == 'fig':
        for fig_set in data:
            for fe in fig_set:
                ce = deepcopy(fe)
                current.addnext(ce)
                current = ce
    elif item_type == 'table' and data is not None:
        ce = deepcopy(data)
        current.addnext(ce)
        current = ce

# ====== Replace placeholders with REF fields ======
print("\nReplacing REF markers in body text...")
marker_map = {
    '【REF1】': 0,
    '【REF2】': 1,
    '【REF3】': 2,
    '【REF4】': 3,
    '【REF5】': 4,
    '【REF6】': 5,
}

for p in doc.paragraphs:
    for marker, ref_idx in marker_map.items():
        if marker in p.text:
            success = replace_marker_in_para(p, marker, xref_paras[ref_idx])
            if success:
                print(f"  ✓ {marker}")

# ====== Save ======
doc.save(target)
print(f"\n✅ Done! Saved to {target}")
