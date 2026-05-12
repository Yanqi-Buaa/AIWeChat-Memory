import docx

DST = r'A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\20260429\测试大纲\11.4-项目技术指标测试报告-北航-20260512-修改版-已修改.docx'
doc = docx.Document(DST)

with open(r'C:\Users\Administrator\cow\tmp\full_text.txt', 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t:
            is_red = False
            for r in p.runs:
                try:
                    if r.font.color and r.font.color.rgb and str(r.font.color.rgb) == 'FF0000':
                        is_red = True
                        break
                except:
                    pass
            tag = ' [RED]' if is_red else ''
            f.write(f'P{i}{tag}: {t}\n\n')

    for ti, table in enumerate(doc.tables):
        f.write(f'\n=== TABLE {ti} ===\n')
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            f.write(f'  Row{ri}: {" | ".join(cells)}\n')

print('done')
