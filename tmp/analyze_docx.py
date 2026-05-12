import docx
import sys

doc = docx.Document(r'A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\20260429\测试大纲\11.4-项目技术指标测试报告-北航-20260512-修改版.docx')

print(f"Total paragraphs: {len(doc.paragraphs)}", flush=True)
print(f"Total tables: {len(doc.tables)}", flush=True)

# Search for modification markers
for i, p in enumerate(doc.paragraphs):
    t = p.text
    if '修改' in t:
        print(f"\n=== Paragraph {i} (len={len(t)}) ===", flush=True)
        # find the position of 修改
        idx = t.find('修改')
        start = max(0, idx - 80)
        end = min(len(t), idx + 500)
        print(t[start:end], flush=True)

# Also dump all non-empty paragraphs to see structure
print("\n\n=== ALL NON-EMPTY PARAGRAPHS ===", flush=True)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t:
        print(f"P{i} [{len(t)} chars]: {t[:120]}...", flush=True)
