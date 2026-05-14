import docx
from lxml import etree

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

print(f"段落总数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

# Check key content
print("\n=== 关键内容检查 ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    style = p.style.name
    if txt and ('概述' in txt[:5] or '成果' in txt[:5] or '结论' in txt[:7] or 
                '图1' in txt[:4] or '图2' in txt[:4] or '图3' in txt[:4] or 
                '图4' in txt[:4] or '图5' in txt[:4] or '图6' in txt[:4] or
                '本项' in txt[:4] or '如' in txt[:10] or '该项' in txt[:4] or
                '交叉' in txt[:4] or '附件' in txt[:4]):
        print(f"[{i:2d}] {style:15s} | {txt[:120]}")

# Count images and field codes
body = doc.element.body
drawings = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing', 
    {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
# Count all drawing elements
from lxml import etree as ET
all_drawings = body.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline')
all_picts = body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict')
print(f"\n图片数量: inline={len(all_drawings)}, pict={len(all_picts)}")

# Count SEQ fields
seq_count = len(body.findall(f'.//{{{ns_w}}}instrText[contains(text(),"SEQ")]'))
ref_count = len(body.findall(f'.//{{{ns_w}}}instrText[contains(text(),"REF")]'))
print(f"SEQ域: {seq_count}, REF域: {ref_count}")

# Check table unchanged
print("\n=== 第1页表格 ===")
for ri, row in enumerate(doc.tables[0].rows):
    cells = [cell.text.strip()[:50] for cell in row.cells]
    print(f"  Row {ri}: {' | '.join(cells)}")

# Check for end content (should be gone)
end_content = sum(1 for p in doc.paragraphs if '交叉引用' in p.text or '图1' in p.text[:4])
print(f"\n末尾残留段落数: {end_content} (应为0)")

# Total word count
total = sum(len(p.text) for p in doc.paragraphs)
print(f"\n文档总字数: {total}")
