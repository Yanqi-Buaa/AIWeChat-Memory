import docx
from lxml import etree

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

print(f"段落总数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

# Count images
body = doc.element.body
all_inline = body.findall(f'.//{{{ns_wp}}}inline')
print(f"图片总数: {len(all_inline)}")

# Count field codes
all_instr = body.findall(f'.//{{{ns_w}}}instrText')
seq_count = sum(1 for s in all_instr if s.text and 'SEQ' in s.text)
ref_count = sum(1 for s in all_instr if s.text and 'REF' in s.text)
print(f"SEQ域: {seq_count}, REF域: {ref_count}")

# Check structure - focus on正文
print("\n=== 正文结构（关键段落）===")
for i, p in enumerate(doc.paragraphs):
    if i < 29:
        continue
    txt = p.text.strip()
    style = p.style.name
    if not txt:
        continue
    # Show first 90 chars
    elem = p._element
    imgs = elem.findall(f'.//{{{ns_wp}}}inline')
    img_flag = ' 🖼️' if imgs else ''
    # Check if contains REF field code inline
    ref_inline = len(elem.findall(f'.//{{{ns_w}}}instrText[contains(text(),"REF")]'))
    seq_inline = len(elem.findall(f'.//{{{ns_w}}}instrText[contains(text(),"SEQ")]'))
    field_flag = ''
    if ref_inline: field_flag += ' [REF]'
    if seq_inline: field_flag += ' [SEQ]'
    
    print(f"[{i:2d}] {style:15s}{img_flag}{field_flag} | {txt[:100]}")

# Check no markers left
markers_left = sum(1 for p in doc.paragraphs if '【REF' in p.text or '【FIG' in p.text or '【XREF' in p.text)
end_left = sum(1 for p in doc.paragraphs if '交叉引用' in p.text)
print(f"\n残留标记: {markers_left} | 末尾残留: {end_left}")

# Check table
print("\n=== 第1页表格 ===")
for ri, row in enumerate(doc.tables[0].rows):
    cells = [cell.text.strip()[:40] for cell in row.cells]
    print(f"  Row {ri}: {' | '.join(cells)}")

if len(doc.tables) > 1:
    print("\n=== 子图表格 ===")
    for ri, row in enumerate(doc.tables[1].rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f"  [{ri},{ci}] {txt}")

# Word count
body_text = sum(len(p.text) for p in doc.paragraphs[29:])
print(f"\n正文字数: {body_text}")
print("\n✅ 验证完成")
