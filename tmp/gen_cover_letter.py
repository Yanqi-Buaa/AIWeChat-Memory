import docx
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def add_run(para, text, bold=False, underline=False, size=12):
    r = para.add_run(text)
    r.bold = bold
    r.underline = underline
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    return r

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Cover Letter', bold=True, size=16)

# Date
p = doc.add_paragraph()
add_run(p, 'Date: May 12, 2026', size=11)

# To
p = doc.add_paragraph()
add_run(p, 'To:', size=11)
p = doc.add_paragraph()
add_run(p, 'Visa Section', size=11)
p = doc.add_paragraph()
add_run(p, 'Consulate General of Italy in Beijing', size=11)

# Applicant
doc.add_paragraph()
p = doc.add_paragraph()
add_run(p, 'Applicant: Qi Yan', size=11)
p = doc.add_paragraph()
add_run(p, 'Passport No.: E94650139', size=11)

# Salutation
doc.add_paragraph()
p = doc.add_paragraph()
add_run(p, 'Dear Visa Officer,', size=12)
doc.add_paragraph()

HL = True  # highlight = bold + underline

body = [
    # Para 1: Background + visa type change
    [
        ("I am writing to re-apply for a Schengen visa following a refusal on April 23, 2026 "
         "(my initial application was submitted on April 17). The refusal cited concerns under "
         "Articles 10, 12, and 13 — regarding the reliability of my stated travel purpose, the "
         "credibility of supporting documents, and my intention to leave the Schengen area before "
         "visa expiry. I believe these concerns arose from a set of materials that were, in hindsight, "
         "incomplete and not fully updated. I have since taken the time to thoroughly review every "
         "document and, more importantly, to understand what went wrong. ", False, False),
        ("In addition, my previous application was submitted under the \"tourism\" category; "
         "upon reflection, \"cultural\" is a more accurate description, as the primary purpose "
         "of my trip is to attend and present at an international academic conference.", False, False),
    ],

    # Para 2: Trip importance — highlight "graduation requirement" and "oral presentation"
    [
        ("This trip is not a casual visit. I am a doctoral candidate at Beihang University, and "
         "attending the ASME 2026 Turbo Expo in Milan (June 15–19, arriving June 14) is a ", False, False),
        ("graduation requirement", HL, HL),
        (". My paper has been accepted for ", False, False),
        ("oral presentation", HL, HL),
        (". If I cannot attend, I may not be able to complete my degree on schedule. This is "
         "the reason I am applying again with great care.", False, False),
    ],

    # Para 3: Updated materials overview
    [
        ("Compared to the previous submission, the following materials have been ", False, False),
        ("newly added or substantially updated", HL, HL),
        (" (see the attached checklist for a full comparison — items in bold indicate changes): "
         "the refusal letter itself (Item 0), financial supporting documents including my parent's "
         "bank statements and securities account records (Item 8), the legal person certificate of "
         "Beihang University with translation (Item 10), the official approval document for my trip "
         "(Item 11), the funding guarantee letter confirming full sponsorship by the university "
         "(Item 12), a detailed day-by-day itinerary (Item 15), the official conference program "
         "showing my session and presentation slot (Item 17), the acceptance letter and the first "
         "page of my paper as proof of authorship (Items 18 & 19), the formal invitation letter "
         "from the conference organizer (Item 20), the registration confirmation and payment invoice "
         "(Item 21), and this updated cover letter (Item 22).", False, False),
    ],

    # Para 4: Session page error
    [
        ("I would also like to clarify one issue that may cause confusion. On the conference "
         "website's session page, the \"Presenting Author\" field mistakenly displays \"Liu Yu\" — "
         "however, the actual presenter introduction on the same page clearly states my name, Qi Yan. "
         "Furthermore, I am the second author of this paper, which independently confirms my role in "
         "presenting it. This is simply a ", False, False),
        ("website data error", HL, HL),
        (", not a discrepancy in my application.", False, False),
    ],

    # Para 5: Hukou
    [
        ("Another point worth clarifying is that my household register (hukou) still lists my "
         "education level as \"high school.\" This is because the register ", False, False),
        ("has not been updated in years", HL, HL),
        (" — a common situation in China. My student ID and the university enrollment certificate, "
         "both included in this submission, clearly establish that I am currently a doctoral student.",
         False, False),
    ],

    # Para 6: Finances + ASME invitation letter clarification
    [
        ("Regarding finances, I have now supplemented my personal bank statements with those of "
         "my parents, as well as detailed securities account records from my father's brokerage. "
         "My father is an active stock investor, and the majority of the family's liquid assets are "
         "held in securities rather than bank deposits — hence the need to present both types of "
         "statements. There is a single large transaction of approximately ", False, False),
        ("RMB 4,500,000", HL, HL),
        (" reflected in the records; this was a ", False, False),
        ("routine wealth management redemption and reallocation", HL, HL),
        (", not an irregular movement of funds. Together, my personal savings, my parents' "
         "financial resources, and the full sponsorship from Beihang University provide more than "
         "adequate coverage for accommodation, meals, and return travel during my stay in Italy. "
         "I note that the ASME invitation letter states that attendees are expected to undertake "
         "their own expenses — this is standard language used by the conference organizer for all "
         "participants. As a student, my travel and living costs are in fact fully covered by "
         "Beihang University, as confirmed by the attached funding guarantee letter (Item 12).",
         False, False),
    ],

    # Para 7: Travel companions
    [
        ("I will be traveling with two colleagues: Dr. Zhuoluo Han (Passport No. EP4961145) and "
         "Researcher Yongfeng Wang (traveling on an official passport). Our group itinerary and "
         "accommodations are fully arranged and documented in the submitted materials.", False, False),
    ],

    # Para 8: Travel history + ties
    [
        ("Finally, I want to underscore two points that I believe are important. First, I have "
         "previously traveled to ", False, False),
        ("Germany, Switzerland, and Austria", HL, HL),
        (" within the Schengen area and ", False, False),
        ("returned to China on time on every occasion", HL, HL),
        (". A copy of my previous passport with the relevant visas and entry/exit stamps is "
         "attached as part of Item 22. Second, I have lived in Beijing for many years and am "
         "deeply committed to completing my doctoral program here. My ongoing research projects, "
         "my academic supervision duties, and my family all bind me firmly to China. I have "
         "every reason — and every intention — to return promptly after the conference.", False, False),
    ],

    # Para 9: Closing
    [
        ("Thank you very much for your time and for reconsidering my application. I am confident "
         "that the substantially strengthened documentation addresses all previous concerns, and "
         "I am happy to provide any further information you may need.", False, False),
    ],
]

for para_segments in body:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75)
    for text, bold, underline in para_segments:
        add_run(p, text, bold=bold, underline=underline, size=12)

# Sign off
doc.add_paragraph()
p = doc.add_paragraph()
add_run(p, 'Sincerely,', size=12)
doc.add_paragraph()
p = doc.add_paragraph()
add_run(p, 'Qi Yan', size=12)

# === ATTACHMENT: FULL CHECKLIST ===
doc.add_page_break()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, 'Attachment: Complete Document Checklist', bold=True, size=13)

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Note: Items in bold are newly added or substantially updated in this submission.')
r.italic = True
r.font.size = Pt(10)
r.font.name = 'Times New Roman'
doc.add_paragraph()

def set_cell_text(cell, text, bold=False, font_size=10):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'

updated_set = {0, 8, 10, 11, 12, 15, 17, 18, 19, 20, 21, 22}

all_items = [
    ('0',  'Refusal Letter',                  '—',               'Newly added'),
    ('1',  'Appointment Confirmation',         'Visa Appointment', 'Submitted previously'),
    ('2',  'Schengen Visa Application Form',   'Visa Appointment', 'Submitted previously'),
    ('3',  'Photo (35×45 mm)',                 'Personal Info',    'Submitted previously'),
    ('4',  'ID Card',                          'Personal Info',    'Submitted previously'),
    ('5',  'Passport + 2 Copies',              'Personal Info',    'Submitted previously'),
    ('6',  'Student ID',                       'Personal Info',    'Re-scanned; submitted previously'),
    ('7',  'Household Register (Hukou) Copy',  'Personal Info',    'Submitted previously'),
    ('8',  'Financial Proof',                  'Personal Info',    'Updated — added parent bank statements & securities records'),
    ('9',  'Enrollment Certificate',           'University Info',  'Submitted previously'),
    ('10', 'Legal Person Certificate',         'University Info',  'Updated — original + translation; approval doc copy added'),
    ('11', 'Official Approval (批件)',          'University Info',  'Newly added'),
    ('12', 'Funding Guarantee Letter',         'University Info',  'Updated — submitted original with full sponsorship'),
    ('13', 'Flight Reservation',               'Travel Info',      'Submitted previously'),
    ('14', 'Accommodation Confirmation',       'Travel Info',      'Submitted previously'),
    ('15', 'Itinerary',                        'Travel Info',      'Updated — refined with detailed daily plan'),
    ('16', 'Insurance Policy (First 3 Pages)', 'Travel Info',      'Submitted previously'),
    ('17', 'Conference Program',               'Travel Info',      'Updated — ASME homepage + session page'),
    ('18', 'Acceptance Proof',                 'Travel Info',      'Updated — acceptance email + paper first page'),
    ('19', 'Paper (First Page)',               'Travel Info',      'Updated — first page of accepted paper'),
    ('20', 'Invitation Letter',                'Travel Info',      'Updated — printed formal invitation'),
    ('21', 'Registration Confirmation',        'Travel Info',      'Updated — confirmation email + invoice'),
    ('22', 'Cover Letter',                     'Travel Info',      'Updated — includes past Schengen travel, companion info, clarifications'),
]

table = doc.add_table(rows=len(all_items) + 1, cols=4, style='Table Grid')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = True

headers = ['No.', 'Item', 'Category', 'Status']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    set_cell_text(cell, h, bold=True, font_size=10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'D9E2F3'
    })
    shading.append(shd)

for row_idx, (no, item, cat, status) in enumerate(all_items):
    is_bold = int(no) in updated_set if no.isdigit() else False
    row_data = [no, item, cat, status]
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        set_cell_text(cell, cell_text, bold=is_bold, font_size=10)
        if col_idx == 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for row in table.rows:
    row.cells[0].width = Cm(1.0)
    row.cells[1].width = Cm(4.5)
    row.cells[2].width = Cm(3.0)
    row.cells[3].width = Cm(7.7)

output_dir = r'A:\0-学习-2025.2-\2-paper\10-2025-ASME GT Turbo 2026\签证\02-文化签'
os.makedirs(output_dir, exist_ok=True)
output_path = os.path.join(output_dir, 'Cover Letter.docx')
doc.save(output_path)
print(f'Saved to: {output_path}')
