# -*- coding: utf-8 -*-
"""
Build the document using ONLY python-docx XML manipulation.
1. Copy backup as working file
2. Write body text
3. Clone figure captions (with SEQ fields), images, cross-refs (with REF fields) from end
4. Insert at correct positions in body
5. Delete end section
6. Keep table 0 (page 1) and table 1 (sub-figures) unchanged
"""
import shutil, os
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paths
backup = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明 - 副本.docx"
target = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"

# Step 0: Copy backup to target
shutil.copy2(backup, target)
doc = Document(target)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

# Helper: create a plain text paragraph element
def make_para(style_name, text):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.append(pStyle)
    p.append(pPr)
    if text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
    return p

def make_para_with_runs(style_name, segments):
    """Create paragraph with multiple runs. segments = [ (text, bold), ... ]"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    pStyle.set(qn('w:val'), style_name)
    pPr.append(pStyle)
    p.append(pPr)
    for text, bold in segments:
        r = OxmlElement('w:r')
        if bold:
            rPr = OxmlElement('w:rPr')
            b = OxmlElement('w:b')
            rPr.append(b)
            r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
    return p

def insert_after(ref_para, elements):
    """Insert list of OxmlElements after ref_para.
    ref_para can be a paragraph or a CT_P element.
    Returns the last inserted element."""
    if hasattr(ref_para, '_element'):
        current = ref_para._element
    else:
        current = ref_para
    for elem in elements:
        current.addnext(elem)
        current = elem
    return current

def clone_element(elem):
    """Deep copy an XML element."""
    return deepcopy(elem)

# ====== Collect end section elements to clone ======
# Cross-references [49]-[54] (REF fields)
xref_elements = [doc.paragraphs[i]._element for i in range(49, 55)]
# Figure captions + images
# [55] image 1
cap_img_1 = [doc.paragraphs[55]._element, doc.paragraphs[56]._element]  # image + caption
# [57] is empty, [58] image 2
cap_img_2 = [doc.paragraphs[58]._element, doc.paragraphs[59]._element]
# [60] empty, [61] image 3
cap_img_3 = [doc.paragraphs[61]._element, doc.paragraphs[62]._element]
# [63] caption 4, [64] image 4
cap_img_4 = [doc.paragraphs[63]._element, doc.paragraphs[64]._element]
# [65] caption 5
cap_img_5 = [doc.paragraphs[65]._element]
# [66][67] empty, [68] image 6
cap_img_6 = [doc.paragraphs[68]._element, doc.paragraphs[69]._element]

# ====== BODY TEXT ======
print("Writing body text...")

# --- 概述 [30] ---
p30 = doc.paragraphs[30]
p30.clear()
p30.add_run(
    '本项目紧密围绕航空发动机整机振动突出、耦合因素多、机理复杂等工程难题，'
    '以某高推重比涡扇发动机为应用对象，开展复杂机动飞行状态下航空发动机转子-支承-机匣系统结构动力学特性研究。'
    '项目突破了气动、机动惯性等复杂载荷作用下发动机结构动力学建模和分析技术，'
    '系统掌握了整机耦合振动机理和动载荷传递规律，探明了机匣连接结构、主轴承及支承构件界面等非线性因素对'
    '转子-支承-机匣系统动力学特性的影响机理和规律，建立了相应的整机结构动力学设计理论与方法，'
    '形成了三项核心成果：'
)
# Add sub-paragraphs for the three achievements
insert_after(p30, [
    make_para('0-报告正文',
        '（1）复杂飞行条件下整机结构系统动力特性仿真与评估技术；'),
    make_para('0-报告正文',
        '（2）基于飞-发多源数据融合的整机振动特征提取与状态识别技术；'),
    make_para('0-报告正文',
        '（3）不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库。'),
    make_para('0-报告正文',
        '上述成果已成功应用于某高推重比涡扇发动机的研制、试飞及交付全过程，'
        '为整机变形评估与振动数据分析提供了有力支撑，应用效果良好。'),
])
print("  概述 done")

# --- 成果1 [32] ---
p32 = doc.paragraphs[32]
p32.clear()
p32.add_run(
    '该项成果主要针对某型小涵道比涡扇发动机的研制需求，重点围绕其机动飞行环境及复杂的进排气条件，'
    '建立了结构特征等效的整机有限元模型，形成了覆盖不同整机运动状态和工作状态的振动响应仿真流程。'
    '在应用过程中，创新性地考虑了机动惯性载荷和气动载荷对整机结构变形的耦合效应，'
    '提出了支承约束力学特性的等效施加方法，突破了传统仿真未充分考虑变形引致支承特性改变的局限。'
)
# Add the second paragraph with figure reference markers
insert_after(p32, [
    make_para('0-报告正文',
        '依托该技术，全面开展了五大类核心仿真分析工作：整机结构建模及力学特性分析、'
        '机动惯性/气动载荷下整机变形仿真、机动惯性载荷下整机振动响应仿真、'
        '加力-喷管气动激励下振动响应特征分析，以及气流旋转激励下的振动响应分析。'
        '部分仿真分析结果如【XREF1】和【XREF2】所示。通过上述仿真与评估，'
        '掌握了复杂载荷对整机结构变形和振动响应的影响规律，主要包括：'),
    make_para('0-报告正文',
        '（1）明确了复杂飞行状态下的惯性与气动载荷主要通过改变整机变形能分布'
        '（表现为支承不同心、轴承内外环倾斜及转静间隙变化），'
        '进而改变转子支承约束力学特性，最终影响转子运动状态及整机振动响应；'),
    make_para('0-报告正文',
        '（2）揭示了机动惯性载荷作用下转子弯曲变形会产生附加旋转惯性力矩的机理，'
        '特别是对于双转子系统，横向载荷会引起转子间进动相互影响，'
        '导致频谱中出现如f1+f2等组合频率成分及多阶模态振动；'),
    make_para('0-报告正文',
        '（3）阐明了加力燃烧室随机宽频气动激励的作用机理，'
        '即承力结构发生受迫振动并对转子形成基础激励，导致支点处激励差异显著，'
        '使转子运动轨迹呈现出特殊的前端外花瓣、后端内花瓣形态；'),
    make_para('0-报告正文',
        '（4）掌握了在旋转惯性、气流旋转及脉动激励综合作用下的能量传递规律，'
        '转静交互激励会产生丰富的倍频特征及幅值波动，'
        '且振动能量在转子与承力结构间发生转移，'
        '显著增加了中介轴承失效及承力结构振动疲劳损伤的风险。'),
    make_para('0-报告正文',
        '该技术融合了机动惯性与气动载荷影响，相比仅考虑不平衡激励的传统方法，'
        '整机振动幅值预测误差降低36%，经试飞500余小时验证，'
        '成功识别了结构变形危险点及关键参数，提出了改进建议，'
        '有力支撑了后续型号衍生设计。'),
])
# Insert marker for figure positions
fig1_marker = insert_after(p32, [make_para('0-报告正文', '【FIG1_2】')])
# Then the ending text
insert_after(fig1_marker, [make_para('0-报告正文',
    '本技术具有工程实用性，可为在研小涵道比涡扇发动机总体结构设计提供定量计算和评估方法参照，'
    '有助于提高涡扇发动机结构设计效率，实现整机振动响应和结构损伤有效控制，以达成结构完整性设计目标。')])
print("  成果1 done")

# --- 成果2 [34] ---
p34 = doc.paragraphs[34]
p34.clear()
p34.add_run(
    '该项成果主要应用于地面台架、高空台及外场飞行试验全过程的监视与诊断。'
    '某高推重比涡扇发动机是我国第一型完全独立自主研制的涡扇发动机，'
    '其工作包线范围、推重比处于世界领先水平。'
    '在外场试飞过程中，曾出现出厂振动合格的发动机左发整机振动幅值异常增大并逼近限值，'
    '而右发正常的现象，如【XREF3】所示。'
)
insert_after(p34, [
    make_para('0-报告正文',
        '项目组将本项目的研究成果成功应用于该型发动机外场飞行振动数据分析，'
        '开展了基于升维扩息的整机振动特征提取及飞-发参数关联性分析，'
        '精准捕捉到转速、飞行速度及俯仰角是对该台份振动异常影响显著的关键参数。'
        '在此基础上，通过融合地面台架数据与装配平衡数据，准确识别了结构状态，'
        '明确指出飞行过程中整机结构状态较地面试车未发生根本性改变，'
        '振动异常增大系装配环节涡轮后支点不同心在飞行气动与机动惯性载荷作用下被放大，'
        '导致振动传递特性改变所致。据此判定整机结构处于第二等级，'
        '虽振幅增加但无明显振动损伤特征，建议后续重点关注整机振动特征及变化趋势，'
        '如【XREF4】和【XREF5】所示。'),
])
fig2_marker = insert_after(p34, [make_para('0-报告正文', '【FIG3_4_5】')])
insert_after(fig2_marker, [
    make_para('0-报告正文',
        '该技术不仅掌握了基于状态切片的振动数据升维扩息方法和高阶特征张量构建方法，'
        '还实现了跨专业多源数据融合与影响因素解耦，'
        '探明了不同转速、气动热力状态及飞行机动载荷对整机振动响应的影响规律。'
        '特别是在出厂交付前，可有效识别出由制造、装配一致性差异导致的'
        '影响飞行安全的异常状态，并指导装配调整。'
        '目前，该技术已在几台份飞行试验及几十台份地面试车中得到应用，'
        '成功筛选并解决多项潜在问题，确保了交付产品质量的一致性，'
        '避免了因非计划下台导致的任务推迟，'
        '为试验考核与分解调整提供了关键决策支持。'),
    make_para('0-报告正文',
        '该技术体系具备良好的通用性与可扩展性。'
        '针对高转速、高负荷条件下的高推重比小涵道比涡扇发动机，'
        '能有效提取复杂工况下的关键结构状态特征，支撑排故分析及视情维护决策。'
        '对于结构相对稳健、载荷环境相对平稳的民用大涵道比涡扇发动机及直升机涡轴发动机，'
        '通过对特征维度进行简化与剪裁，同样适用其状态监测需求。'),
])
print("  成果2 done")

# --- 成果3 [36] ---
p36 = doc.paragraphs[36]
p36.clear()
p36.add_run(
    '该项成果采用短时傅里叶变换实现高频采样信号的频域转化与数据压缩，'
    '提取转速基频、模态频率及频域能量分布特征。'
    '基于成果中的整机特征提取与分析技术，利用四阶张量对发动机结构状态特征进行多维描述'
    '（架次-切片-参数-统计量）。'
    '通过对多台同型号发动机历史数据的集成，'
    '构建了\u201c不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库\u201d。'
    '该数据库基于实测数据提炼关键特征参数，直接服务于飞行载荷确定及振动基准建立，'
    '如【XREF6】所示。随着数据持续积累，为服役后的异常边界判定与精细化健康管理奠定了基础。'
)
fig3_marker = insert_after(p36, [make_para('0-报告正文', '【FIG6】')])
insert_after(fig3_marker, [
    make_para('0-报告正文',
        '该数据库主要应用于全寿命周期健康管理与标准体系建设：'
        '通过持续积累并提炼典型飞行状态下的振动特征参数，'
        '建立了实际服役环境下的整机振动水平基准与量化评估标准，'
        '不仅为批产发动机的制造装配质量评估与异常边界精确判定提供了数据资产，'
        '还可反馈指导复杂载荷环境下发动机的结构健康管理策略制定与改进设计，'
        '提升机群整体的运行安全性与任务可靠性。'),
])
print("  成果3 done")

# --- 结论 [38] ---
p38 = doc.paragraphs[38]
p38.clear()
p38.add_run(
    '本项目形成的\u201c复杂机动飞行状态下航空发动机整机动力特性评估、结构设计与监视诊断技术\u201d，'
    '从仿真评估、状态监视识别到数据基准构建，'
    '形成了一套自主可控的面向真实服役环境的航空发动机整机结构动力特性评估与状态识别体系。'
)
insert_after(p38, [
    make_para('0-报告正文',
        '三项核心成果均已成功应用于某高推重比涡扇发动机的研制全过程：'),
    make_para('0-报告正文',
        '成果一（复杂飞行条件下整机结构系统动力特性仿真与评估技术）'
        '服务于研发设计与验证阶段，支撑了结构薄弱环节识别与整机动力学设计优化；'),
    make_para('0-报告正文',
        '成果二（基于飞-发多源数据融合的整机振动特征提取与状态识别技术）'
        '应用于地面台架与外场试飞全过程的监视诊断，有效识别多起潜在异常并指导装配调整；'),
    make_para('0-报告正文',
        '成果三（不同状态下高推重比发动机及大功率涡桨发动机典型振动特征参数数据库）'
        '服务于全寿命周期健康管理，建立了整机振动水平基准与量化评估标准。'),
    make_para('0-报告正文',
        '综上所述，本项目成果在型号研制与生产中应用效果良好，'
        '有效提升了复杂服役环境下航空发动机整机结构动力学评估与状态诊断能力，'
        '保障了装备的研制进度与运行安全，具有广阔的工程应用前景。'),
])
print("  结论 done")

# ====== INSERT FIGURES & CROSS-REFS ======
print("\nInserting figures and cross-references...")

# Find the marker paragraphs (they've shifted index due to insertions)
# Search by text content
all_paras = doc.paragraphs

def find_para_by_text(substring):
    for p in all_paras:
        if substring in p.text:
            return p
    return None

# Clone cross-references (replace markers)
xref_map = [
    ('【XREF1】', xref_elements[0]),  # 图1 REF
    ('【XREF2】', xref_elements[1]),  # 图2 REF
    ('【XREF3】', xref_elements[2]),  # 图3 REF
    ('【XREF4】', xref_elements[3]),  # 图4 REF
    ('【XREF5】', xref_elements[4]),  # 图5 REF
    ('【XREF6】', xref_elements[5]),  # 图6 REF
]

for marker, xref_elem in xref_map:
    p = find_para_by_text(marker)
    if p:
        # Replace marker text with cloned cross-ref element
        clone = clone_element(xref_elem)
        p._element.addnext(clone)
        # Remove the marker text from the paragraph
        txt = p.text
        p.clear()
        # Re-add text without marker
        cleaned = txt.replace(marker, '')
        if cleaned:
            p.add_run(cleaned)
        print(f"  Replaced {marker}")

# Insert figure captions + images
fig_map = [
    ('【FIG1_2】', cap_img_1 + cap_img_2),  # Fig1 caption+img, Fig2 caption+img
    ('【FIG3_4_5】', cap_img_3 + cap_img_4 + cap_img_5),  # Fig3, Fig4, Fig5
    ('【FIG6】', cap_img_6),  # Fig6 caption+img
]

for marker, fig_elems in fig_map:
    p = find_para_by_text(marker)
    if p:
        ref_elem = p._element
        for elem in fig_elems:
            clone = clone_element(elem)
            ref_elem.addnext(clone)
            ref_elem = clone
        # Remove marker paragraph
        p._element.getparent().remove(p._element)
        print(f"  Inserted figures at {marker}")

# ====== DELETE END SECTION ======
print("\nDeleting end section...")
# Paragraphs 48-70 currently contain the end content
# But indices may have shifted. Let me find and delete by content
end_markers = ['交叉引用：', '图1', '图2', '图3', '图4', '图5', '图6',
               '横向惯性载荷作用下', '弹性线示意图', '振动异常', '特征提取与分析',
               '应用效果', '数据库示意图']

# Find all paragraphs that are AFTER the正文 headings and contain end content
# Better approach: Find the "交叉引用：" paragraph and delete from there to end
xref_header = find_para_by_text('交叉引用：')
if xref_header:
    # Delete xref_header and all following siblings
    parent = xref_header._element.getparent()
    # Collect elements to delete
    to_delete = []
    next_elem = xref_header._element
    while next_elem is not None:
        to_delete.append(next_elem)
        next_elem = next_elem.getnext()
    # Delete
    for elem in to_delete:
        parent.remove(elem)
    print(f"  Deleted {len(to_delete)} elements from end section")

# ====== SAVE ======
doc.save(target)
print(f"\n✅ 文档已保存至: {target}")
print("=" * 50)
print("文档结构：")
print("  一、概述")
print("  二、成果1的应用情况及效果（含图1、图2及对应交叉引用）")
print("  三、成果2的应用情况及效果（含图3、图4、图5及对应交叉引用）")
print("  四、成果3的应用情况及效果（含图6及对应交叉引用）")
print("  五、成果应用结论")
print("=" * 50)
