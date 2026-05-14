import docx
from lxml import etree

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

print(f"段落总数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

# Count images
body = doc.element.body
all_inline = body.findall(f'.//{{{ns_wp}}}inline')
print(f"图片总数: {len(all_inline)}")

# Count SEQ and REF
all_instr = body.findall(f'.//{{{ns_w}}}instrText')
seq_count = sum(1 for s in all_instr if s.text and 'SEQ' in s.text)
ref_count = sum(1 for s in all_instr if s.text and 'REF' in s.text)
print(f"SEQ域: {seq_count}, REF域: {ref_count}")

# Check for cross-ref fields - the REF fields come in pairs (begin+separate+end)
fld_chars = body.findall(f'.//{{{ns_w}}}fldChar')
print(f"域代码字符总数: {len(fld_chars)}")

# Print structure
print("\n=== 文档正文结构 ===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    style = p.style.name
    if txt:
        elem = p._element
        imgs = elem.findall(f'.//{{{ns_wp}}}inline')
        img_flag = ' 🖼️' if imgs else ''
        if i < 29:
            if i < 25:  # cover area
                if i in [0,1,2,7,8,13,14,15,16,17,18,19,23]:
                    print(f"[{i:2d}] {style:15s} | {txt[:80]}")
        else:
            print(f"[{i:2d}] {style:15s}{img_flag} | {txt[:110]}")

# Check table 0
print("\n=== 第1页表格 ===")
for ri, row in enumerate(doc.tables[0].rows):
    cells = [cell.text.strip()[:50] for cell in row.cells]
    print(f"  Row {ri}: {' | '.join(cells)}")

# Check for extra table (sub-figure)
if len(doc.tables) > 1:
    print("\n=== 第2个表格（子图）===")
    for ri, row in enumerate(doc.tables[1].rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f"  [{ri},{ci}] {txt[:80]}")

# Count total body text
body_text = sum(len(p.text) for p in doc.paragraphs[29:])
print(f"\n正文字数: {body_text}")

# Check no残留 markers
markers_remaining = sum(1 for p in doc.paragraphs if '【' in p.text or 'XREF' in p.text or 'FIG' in p.text)
print(f"残留标记: {markers_remaining} (应为0)")

# Check no end-section leftover
end_left = sum(1 for p in doc.paragraphs if '交叉引用' in p.text)
print(f"末尾残留: {end_left} (应为0)")

print("\n✅ 验证完成")
