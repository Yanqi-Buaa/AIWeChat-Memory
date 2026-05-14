import docx, os

base = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明"

files = [
    "（公开）附件1：应用证明需求申请表-20260203.wps",
    "（公开）附件3：成果应用情况审查报告-20260203.wps", 
    "（公开）附件4：应用证明模板-20260203.wps"
]

for fname in files:
    path = os.path.join(base, fname)
    size = os.path.getsize(path)
    print(f"\n{'='*60}")
    print(f"FILE: {fname} ({size} bytes)")
    print('='*60)
    try:
        doc = docx.Document(path)
        print(f"Paragraphs: {len(doc.paragraphs)}, Tables: {len(doc.tables)}")
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip():
                print(f"[{i}] {p.style.name} | {p.text}")
        if doc.tables:
            for ti, table in enumerate(doc.tables):
                print(f"\n--- Table {ti} ({len(table.rows)} x {len(table.columns)}) ---")
                for ri, row in enumerate(table.rows):
                    cells = [cell.text.strip()[:50] for cell in row.cells]
                    print(f"  Row {ri}: {' | '.join(cells)}")
    except Exception as e:
        print(f"ERROR: {e}")
        # Try as plain text
        try:
            with open(path, 'r', encoding='utf-8') as f:
                print(f.read()[:2000])
        except:
            with open(path, 'rb') as f:
                header = f.read(100)
                print(f"Binary header (hex): {header.hex()}")
