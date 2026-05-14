import docx
from lxml import etree

path = r"A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\应用证明\应用评价证明\基础研究成果应用评价证明.docx"
doc = docx.Document(path)

ns_w = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
ns_wp = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'

print(f"段落总数: {len(doc.paragraphs)}")
print(f"表格数: {len(doc.tables)}")

# Count images
body = doc.element.body
all_inline = body.findall(f'.//{{{ns_wp}}}inline')
print(f"图片总数: {len(all_inline)}")

# Count field codes (iterate through all instrText)
all_instr = body.findall(f'.//{{{ns_w}}}instrText')
seq_count = 0
ref_count = 0
for s in all_instr:
    if s.text and 'SEQ' in s.text:
        seq_count += 1
    if s.text and 'REF' in s.text:
        ref_count += 1
print(f"SEQ域: {seq_count}, REF域: {ref_count}")

# Count fldChars
all_fld = body.findall(f'.//{{{ns_w}}}fldChar')
print(f"域代码字符: {len(all_fld)}")

# Show structure
print("\n=== 正文结构 ===")
for i, p in enumerate(doc.paragraphs):
    if i < 29:
        continue
    txt = p.text.strip()
    style = p.style.name
    if not txt:
        continue
    
    elem = p._element
    imgs = len(elem.findall(f'.//{{{ns_wp}}}inline'))
    
    # Check for field codes by iterating
    has_ref = False
    has_seq = False
    for instr in elem.findall(f'.//{{{ns_w}}}instrText'):
        if instr.text:
            if 'REF' in instr.text: has_ref = True
            if 'SEQ' in instr.text: has_seq = True
    
    flag = ''
    if imgs: flag += ' 🖼️'
    if has_ref: flag += ' [REF]'
    if has_seq: flag += ' [SEQ]'
    
    print(f"[{i:2d}] {style:15s}{flag} | {txt[:100]}")

# Check markers
markers_left = sum(1 for p in doc.paragraphs if '【REF' in p.text or '【FIG' in p.text or '【XREF' in p.text)
end_left = sum(1 for p in doc.paragraphs if '交叉引用' in p.text)
print(f"\n残留标记: {markers_left} | 末尾残留: {end_left}")

# Check table
print("\n=== 第1页表格 ===")
for ri, row in enumerate(doc.tables[0].rows):
    cells = [cell.text.strip()[:40] for cell in row.cells]
    print(f"  Row {ri}: {' | '.join(cells)}")

if len(doc.tables) > 1:
    print("\n=== 子图表格 ===")
    for ri, row in enumerate(doc.tables[1].rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f"  [{ri},{ci}] {txt}")

# Word count
body_text = sum(len(p.text) for p in doc.paragraphs[29:])
print(f"\n正文字数: {body_text}")
print("\n✅ 验证完成")
