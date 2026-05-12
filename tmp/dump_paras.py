import docx, sys
doc = docx.Document(r'A:\0-学习-2025.2-\A 2023.1 三批lj申请\00-飞行状态\结题\C4.4课题验收\20260429\测试大纲\11.4-项目技术指标测试报告-北航-20260512-修改版.docx')

key_paras = [75,76,77,78,79,80,81,82,83,84,85,86,87,88,89, 98,99,100, 113,114,115,116,117,118,119,120,121,122,123,124]

with open(r'C:\Users\Administrator\cow\tmp\para_dump.txt', 'w', encoding='utf-8') as f:
    for i in key_paras:
        if i < len(doc.paragraphs):
            t = doc.paragraphs[i].text
            f.write(f'=== P{i} (len={len(t)}) ===\n')
            f.write(t + '\n\n')

print('Done', flush=True)
