import docx, os

# Read all paragraphs from attachment 3 (the main content source)
path = r"C:\Users\Administrator\cow\tmp\（公开）附件3：成果应用情况审查报告-20260203.docx"
doc = docx.Document(path)

print("=== 附件3 所有段落（含空行）===")
for i, p in enumerate(doc.paragraphs):
    txt = p.text
    style = p.style.name
    print(f"[{i}] {style} | {repr(txt)}")

print("\n\n=== 附件3 表格 ===")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti} ({len(table.rows)} x {len(table.columns)}) ---")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f"  [{ri},{ci}] {txt[:200]}")

# Check for images
print("\n\n=== 附件3 图片检查 ===")
from docx.opc.constants import RELATIONSHIP_TYPE as RT
rels = doc.part.rels
img_count = 0
for rel_id, rel in rels.items():
    if "image" in str(rel.reltype).lower():
        img_count += 1
        print(f"  Image: {rel.target_ref}, type={rel.reltype}")
print(f"  共 {img_count} 张图片")

# Also check inline shapes
print("\n=== InlineShapes检查 ===")
for i, p in enumerate(doc.paragraphs):
    if p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
        print(f"  [{i}] 有drawing元素: {p.text[:60]}")
    if p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
        print(f"  [{i}] 有pict元素: {p.text[:60]}")

# Also check attachment 1 table content more fully
path1 = r"C:\Users\Administrator\cow\tmp\（公开）附件1：应用证明需求申请表-20260203.docx"
doc1 = docx.Document(path1)
print("\n\n=== 附件1 完整表格 ===")
for ti, table in enumerate(doc1.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f"  [Row{ri},Col{ci}] {txt}")
